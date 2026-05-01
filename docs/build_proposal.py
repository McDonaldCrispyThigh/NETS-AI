"""
Generate Research_Proposal.docx in the docs/ folder.
Run: python docs/build_proposal.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# Base font
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

# ── Helpers ────────────────────────────────────────────────────

def font(run, bold=False, italic=False, size=12, color=None):
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def H(text, center=False, size=14, bold=True, color=None,
      space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, bold=bold, size=size, color=color)
    return p

def B(text, indent=False, space_after=6, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(text)
    font(r, size=size)
    return p

def shade_cell(cell, hex_color='D9D9D9'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        font(r, bold=True, size=10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(c)
    for ri, row_data in enumerate(rows):
        for ci, txt in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(txt)
            font(r, size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def bullet(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label + '  ')
    font(r1, bold=True)
    r2 = p.add_run(text)
    font(r2)

# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HONORS THESIS RESEARCH PROPOSAL')
font(r, bold=True, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run('Can Artificial Intelligence Reconstruct Commercial Geographies?')
font(r, bold=True, size=16, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run(
    'Evaluating AI-Generated Pharmacy Establishment Data as a Validator and '
    'Alternative to the National Establishment Time-Series (NETS) Database '
    'in Minneapolis, Minnesota'
)
font(r, italic=True, size=13)

for lbl, val in [
    ('Student Investigator:', 'Congyuan Zheng'),
    ('Department:',           'Geography, University of Colorado Boulder'),
    ('Primary Supervisor:',   'Prof. Jessica Finlay, CU Boulder'),
    ('Co-Supervisor:',        'Prof. Michael Esposito, University of Minnesota'),
    ('Postdoctoral Mentor:',  'Yue Sun, CU Boulder'),
    ('Submission Date:',      'April 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(lbl + '  ')
    font(r1, bold=True)
    r2 = p.add_run(val)
    font(r2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════
H('Abstract', center=True, size=13)
B(
    'The National Establishment Time-Series (NETS) database is a foundational data '
    'source for urban commercial geography, yet its construction relies on proprietary '
    'Dun and Bradstreet credit records and remains methodologically opaque. A 2017 '
    'Federal Reserve Board assessment confirmed that NETS exhibits its largest coverage '
    'gaps among small establishments and operates on a two-to-three-year lag in '
    'recording business births and deaths (Barnatchez, Crane, and Decker 2017). These '
    'limitations are rarely acknowledged in health geography studies that use NETS to '
    'document pharmacy access, despite a worsening crisis: 8 percent of Minnesota '
    'residents now live in a pharmacy desert (up from 6.2 percent in 2009), and '
    'approximately 34 percent of Twin Cities residents live in or near one (Minnesota '
    'Department of Health 2024). North Minneapolis has experienced disproportionate '
    'chain closures, with Walgreens shuttering its W. Broadway location in February '
    '2023 and CVS following with four additional area closures, leaving predominantly '
    'Black, lower-income residents without nearby pharmacy access.'
)
B(
    'This thesis develops and deploys a portable, reproducible Multi-Agent System '
    '(MAS) that collects, classifies, and structures pharmacy establishment data for '
    'the Minneapolis-St. Paul metropolitan area using the Google Maps Places API, '
    'Yelp Fusion API, and OpenAI GPT-4o-mini (temperature 0.0 for deterministic '
    'outputs), orchestrated via LangChain across four specialized agents: Search, '
    'Enrichment, Classification, and QA. The resulting dataset is compared against '
    'verified ground-truth records from the Minnesota Board of Pharmacy licensure '
    'database and against NETS-derived pharmacy records across three dimensions: '
    'spatial coverage, NAICS classification accuracy, and chain versus independent '
    'business composition. Spatial analysis is conducted in PostGIS with Bayesian '
    'hierarchical regression implemented in PyMC, enabling uncertainty quantification '
    'in coverage gap estimates. The central analytical contribution is triangulation: '
    'where AI and NETS diverge from regulatory ground truth, this study identifies '
    'whose error is larger, in what spatial pattern, and what that pattern reveals '
    'about each dataset\'s structural biases. If NETS systematically undercounts '
    'independent pharmacies in disadvantaged neighborhoods, existing pharmacy desert '
    'research may have understated the severity of the access crisis it claims to measure.'
)
p = doc.add_paragraph()
r1 = p.add_run('Keywords:  ')
font(r1, bold=True)
r2 = p.add_run(
    'NETS Database  |  Pharmacy Deserts  |  Health Equity  |  AI Agents  |  '
    'Minneapolis  |  Reproducibility  |  Urban Commercial Geography  |  GPT-4o-mini'
)
font(r2, italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════
H('1.  Introduction and Background')

H('1.1  The Pharmacy Desert Crisis in Minnesota', size=12)
B(
    'Access to community pharmacies is a critical social determinant of health. '
    'Pharmacies serve as first-contact healthcare providers, dispensing prescription '
    'medications, administering immunizations, and delivering preventive care '
    'consultations. Their spatial distribution is a core dimension of urban health '
    'equity, and its deterioration constitutes a measurable public health emergency.'
)
B(
    'Minnesota is confronting an acute and worsening pharmacy desert crisis. '
    'A 2024 report by the Minnesota Department of Health found that 8 percent of '
    'state residents (approximately 463,000 people) now live in a pharmacy desert, '
    'up from 6.2 percent in 2009. In the Twin Cities metropolitan area the situation '
    'is more severe: approximately 34 percent of residents live in or near a pharmacy '
    'desert. North Minneapolis has experienced especially pronounced deterioration. '
    'Walgreens shuttered its W. Broadway Avenue location in February 2023; CVS '
    'subsequently closed four additional area stores. Residents in these predominantly '
    'Black, lower-income neighborhoods now face multi-mile trips to access the nearest '
    'pharmacy, with only the Cub Pharmacy on West Broadway remaining as a nearby '
    'option (Fox9 Minneapolis 2024; Star Tribune 2024).'
)
B(
    'Statewide, approximately 44 percent of all Minnesota pharmacies have closed over '
    'the past decade, with independent pharmacies accounting for nearly 60 percent of '
    'those closures (NCPA State of the Indy Report 2026; MinnPost 2025). In 2025 and '
    '2026, Minnesota independent pharmacists testified before the state legislature '
    'seeking emergency PBM reimbursement reform. The crisis demands precise, spatially '
    'complete data on pharmacy location and composition. Yet the data infrastructure '
    'that health geography relies upon to measure and map this crisis carries its own '
    'unexamined vulnerabilities.'
)

H('1.2  The Problem: NETS as Standard, NETS as Black Box', size=12)
B(
    'The National Establishment Time-Series (NETS) database, compiled by Walls and '
    'Associates from Dun and Bradstreet (D and B) credit records, provides '
    'longitudinal establishment-level data covering business names, addresses, NAICS '
    'codes, employee counts, and year of establishment dating to the late 1980s. Its '
    'longitudinal depth and spatial granularity have made it the de facto standard for '
    'tracking commercial change in American cities. Pharmacy desert studies, '
    'neighborhood retail dynamics research, and commercial displacement analyses '
    'routinely draw on NETS without interrogating its underlying accuracy.'
)
B(
    'A rigorous 2017 assessment by economists at the Federal Reserve Board and the '
    'Federal Reserve Bank of Minneapolis (Barnatchez, Crane, and Decker 2017) '
    'documented fundamental limitations with direct implications for pharmacy research. '
    'First, NETS exhibits its largest coverage gaps among small establishments, '
    'precisely the independent pharmacies and community-serving businesses most '
    'relevant to health equity research. When NETS is restricted to medium-to-large '
    'employers, it accounts for approximately 73 percent of U.S. employment in the '
    'Quarterly Census of Employment and Wages; for small establishments, reliability '
    'falls substantially. Second, NETS operates on a two-to-three-year lag in '
    'recording business births and deaths: a closed pharmacy may remain in the dataset '
    'for years after closure, and a newly opened one may be absent. Third, D and B '
    'data has documented coverage gaps for minority-owned and women-owned businesses, '
    'confirmed by an independent state-level evaluation in Washington (WA Dept. of '
    'Commerce 2021). Fourth, business dynamics in NETS microdata are markedly '
    'different from official Census administrative sources, limiting NETS usefulness '
    'for studying establishment entry and exit.'
)
B(
    'The health equity implication is direct and underappreciated. Research documenting '
    'that North Minneapolis has fewer pharmacies than wealthier neighborhoods may be '
    'partially correct as empirical fact but partially an artifact of NETS '
    'systematically undercounting the types of pharmacies that serve disadvantaged '
    'communities. Current scholarship cannot distinguish these two explanations. '
    'This thesis provides the analytical instrument to do so.'
)

H('1.3  The Opportunity: AI as Independent Validator', size=12)
B(
    'The convergence of large language models and public business data infrastructure '
    'creates a new methodological possibility. It is now feasible to construct a '
    'reproducible AI agent that queries public APIs to enumerate real-world '
    'establishments, extracts structured attributes from observable signals, uses an '
    'LLM to classify establishments by industry code, and outputs a structured dataset '
    'aligned to the NETS schema at a marginal cost below USD 100 per city. Unlike '
    'NETS, every step of this process is transparent, documented, and replicable.'
)
B(
    'Recent scholarship confirms the feasibility. A 2025 ISPRS study on LLMs for '
    'GIS-based spatial analysis finds competitive classification accuracy when '
    'contextual prompts are well-designed. A 2025 systematic review in the Journal '
    'of Urban Technology identifies structured entity classification and attribute '
    'extraction as demonstrated LLM capabilities in urban contexts. The Google Maps '
    'Places API covers over 200 million places globally and has served as a primary '
    'data source in peer-reviewed research on restaurant survival and commercial '
    'geography. A 2025 ScienceDirect study on restaurant survival using nationwide '
    'Google Maps data established that the platform, when cleaned of inauthentic '
    'listings, produces research-grade commercial geography data.'
)
B(
    'Critically, the AI agent need not be perfect to be useful. Its errors are '
    'informative. By comparing AI outputs, NETS records, and verified regulatory '
    'ground truth from the Minnesota Board of Pharmacy simultaneously, this study '
    'identifies not only where AI falls short, but where NETS has been providing an '
    'incomplete picture of commercial geography. The triangulation is the '
    'methodological contribution.'
)

# ════════════════════════════════════════════════════════════════
# 2. RESEARCH QUESTIONS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
H('2.  Research Questions and Hypotheses')

B('This study addresses four primary research questions:')
bullet('RQ1 (Coverage):',
       'To what extent does the AI-generated pharmacy dataset capture the complete '
       'population of active Minneapolis pharmacies, as verified against Minnesota '
       'Board of Pharmacy licensure records?')
bullet('RQ2 (Accuracy):',
       'How accurately does the AI agent classify pharmacies by NAICS code 446110, '
       'and what establishment types generate systematic misclassification?')
bullet('RQ3 (Divergence):',
       'Where the AI dataset and NETS diverge from regulatory ground truth, do their '
       'errors follow distinct spatial or organizational patterns, and what do those '
       'patterns reveal about each dataset\'s structural biases?')
bullet('RQ4 (Equity Implications):',
       'Are identified biases spatially correlated with neighborhood-level '
       'socioeconomic disadvantage, and if so, do they systematically understate '
       'pharmacy desert severity in Minneapolis minority and low-income communities?')

B('')
B(
    'Primary hypothesis: NETS will exhibit lower recall relative to Minnesota Board '
    'of Pharmacy ground truth than the AI agent, specifically among independent '
    'pharmacies and in low-income ZIP codes. This directional bias, if confirmed, '
    'suggests that pharmacy desert analyses built on NETS have understated the '
    'severity of access gaps in disadvantaged neighborhoods by undercounting the '
    'pharmacies that serve them. The AI agent, drawing on Google Maps\'s user-contributed '
    'listings, may exhibit a complementary bias: overrepresentation of digitally '
    'visible chain pharmacies relative to independent ones. Characterizing these '
    'divergent error structures is the core analytical task.'
)

# ════════════════════════════════════════════════════════════════
# 3. LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════
H('3.  Literature Review')

H('3.1  Pharmacy Access and Health Equity', size=12)
B(
    'Pharmacy desert research gained national prominence with Qato et al. (2014), '
    'who applied a half-mile buffer criterion in Chicago and found that predominantly '
    'Black and Hispanic communities faced systematically greater barriers to pharmacy '
    'access. This framework has been extended nationally by Dill and Gelmon (2023), '
    'who used a combination of NETS and Google Places data but did not validate either '
    'source against regulatory records. A 2025 study introducing a travel-time-based '
    'metric (reported in Medical Xpress, April 2025) and the USC pharmacy desert '
    'mapping tool relaunch (Daily Trojan, November 2025) indicate active methodological '
    'development in the field, but data source validation remains unaddressed.'
)
B(
    'Minnesota-specific research has documented the crisis with increasing granularity. '
    'The University of Minnesota College of Pharmacy conducted longitudinal pharmacy '
    'desert mapping at five time points from 2009 to 2024, presenting findings at '
    'the 2024 APHA Annual Meeting (Abstract 551580). Their work documents the specific '
    'harm to North Minneapolis following chain closures. A 2025 PLOS ONE study on the '
    'association between pharmacy desert distribution, digital divide, and residential '
    'redlining confirms a structural link between historical disinvestment and current '
    'pharmacy access barriers. What none of these studies examines is whether the '
    'pharmacy enumeration underpinning their analyses is itself complete and unbiased.'
)

H('3.2  NETS: Strengths, Limitations, and the Minneapolis Federal Reserve Assessment', size=12)
B(
    'NETS has served as the primary commercial establishment database in urban '
    'geography for over two decades. Its application spans retail displacement '
    '(Chapple and Jacobus 2009), neighborhood commercial dynamics (Meltzer and '
    'Schuetz 2012), and business survival analysis. The database\'s longitudinal '
    'depth is its primary asset; no other publicly accessible source provides '
    'establishment-level time series at comparable spatial granularity.'
)
B(
    'The Federal Reserve Board assessment (Barnatchez, Crane, and Decker 2017), '
    'independently reviewed by the Federal Reserve Bank of Minneapolis, is the most '
    'comprehensive methodological critique of NETS to date. Its findings are '
    'unambiguous: NETS\'s reliability is highest for medium-to-large establishments '
    'and degrades significantly for small, single-location businesses. NETS misses '
    'key economic dynamics including post-2000 manufacturing decline and post-2007 '
    'construction contraction. A 2019 Finance and Economics Discussion Series paper '
    'further documents discrepancies between NETS and the Census Bureau Longitudinal '
    'Business Database. The Washington State D and B evaluation (2021) confirms '
    'undercoverage of minority-owned businesses in D and B-derived data. These '
    'findings have not significantly altered NETS\'s use in pharmacy access research, '
    'creating the methodological gap this thesis addresses.'
)

H('3.3  AI Agents and LLMs in Urban Geography', size=12)
B(
    'LLM applications in urban planning and geography have expanded rapidly since '
    '2023. A 2025 systematic review in the Journal of Urban Technology documents '
    'the trajectory from early experimentation to structured deployment in urban '
    'analysis tasks. Studies on LLMs for GIS-based spatial analysis (Shen et al. '
    '2025, ISPRS) demonstrate that GPT-4 achieves competitive accuracy in '
    'transportation and geographic classification tasks. Research on generative AI '
    'for urban digital twins establishes conceptual frameworks for AI-generated '
    'geospatial data products. A ScienceDirect machine learning review (2025) '
    'identifies structured entity classification and attribute extraction as the '
    'highest-confidence LLM capabilities in urban contexts.'
)
B(
    'On the data quality side, a 2025 ScienceDirect study predicting restaurant '
    'survival using nationwide Google Maps data demonstrates that the platform, '
    'after filtering inauthentic listings (identified as zero-review or rapidly '
    'delisted entries), produces reliable commercial geography data for academic '
    'research. This validates the methodological approach adopted in this thesis '
    'while confirming the need for systematic quality controls.'
)

H('3.4  Ground Truth: NCPDP DataQ and State Pharmacy Licensure', size=12)
B(
    'The National Council for Prescription Drug Programs (NCPDP) DataQ database '
    'provides verified information on approximately 80,000 U.S. pharmacies, including '
    'NPI numbers, license numbers, Medicare and Medicaid identifiers, 340B status, '
    'and immunization service indicators. Unlike NETS, DataQ is maintained by a team '
    'exclusively focused on pharmacy data verification and updated in real time. '
    'Recent pharmacy desert research, including the national cross-sectional study '
    'presented at the 2024 APHA Annual Meeting, has adopted NCPDP DataQ as the '
    'preferred alternative to NETS for pharmacy enumeration, precisely because of '
    'NETS\'s coverage limitations.'
)
B(
    'The Minnesota Board of Pharmacy maintains a public licensure database, accessible '
    'via formal data request through mn.gov/boards/pharmacy/public/datarequests.jsp, '
    'covering all licensed retail pharmacies with address, license type, active status, '
    'and issue date. This study uses this database as its primary ground truth, '
    'supplemented by NCPDP DataQ, providing a regulatory anchor unavailable in '
    'prior pharmacy desert studies.'
)

# ════════════════════════════════════════════════════════════════
# 4. METHODOLOGY
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
H('4.  Methodology')

H('4.1  Study Area', size=12)
B(
    'The study area is the Minneapolis-St. Paul metropolitan area, defined by '
    'approximately 40 ZIP codes spanning the two core cities: Minneapolis '
    '(55401 through 55415, 55454, and 55455; 17 ZIP codes) and St. Paul '
    '(55101 through 55108, 55116 through 55119, and 55130; approximately 14 ZIP '
    'codes), together covering the full Hennepin and Ramsey county urban cores. '
    'This expanded scope captures the complete Twin Cities pharmacy landscape '
    'documented in recent Minnesota Department of Health reporting, in which the '
    'metropolitan area is treated as a single functional unit for healthcare '
    'access analysis.'
)
B(
    'The metropolitan area is selected for four reasons: (1) it presents a '
    'well-documented and actively worsening pharmacy desert crisis, with specific '
    'closure events verifiable against news records; (2) active pharmacy desert '
    'research at the University of Minnesota provides methodological comparators; '
    '(3) both Minnesota Board of Pharmacy and Ramsey County public licensure '
    'records are accessible; and (4) the demographic composition, including the '
    'predominantly Black communities of North Minneapolis and the immigrant and '
    'refugee communities of East St. Paul, enables equity-stratified analysis '
    'across multiple disadvantaged subpopulations. Estimated total pharmacy '
    'records: 2,500 to 3,500 after deduplication across all ZIP codes. The '
    'temporal scope is a point-in-time cross-section as of 2025 to 2026, with '
    'longitudinal context from NETS time-series records where available.'
)

H('4.2  Data Sources and Computational Stack', size=12)
B('External data sources:')
table(
    headers=['Source', 'Role in Study', 'Access Method'],
    rows=[
        ['Google Maps Places API',    'Primary business search and attribute enrichment',            'Paid API'],
        ['OpenAI GPT-4o-mini',        'NAICS classification; metadata estimation (temp. = 0.0)',     'Paid API'],
        ['Yelp Fusion API',           'Secondary enrichment: open/closed status, categories',        'Free tier'],
        ['MN Board of Pharmacy',      'Primary regulatory ground truth',                             'Public data request'],
        ['NCPDP DataQ',               'Supplementary ground truth validation',                       'Institutional access'],
        ['NETS Database',             'Comparison dataset (via U of Minnesota)',                     'Institutional access'],
        ['U.S. Census ACS 2019-2023', 'Socioeconomic covariates for spatial equity analysis',       'Public API'],
    ],
    widths=[1.8, 2.8, 1.6]
)
B('Computational and analytical tools:')
table(
    headers=['Tool', 'Role', 'Justification'],
    rows=[
        ['LangChain',   'Multi-agent orchestration layer; manages agent handoffs, memory, and tool routing', 'Industry-standard framework for composable LLM agent systems; enables modular MAS design'],
        ['PostGIS',     'Spatial data storage, indexing, and SQL-based spatial queries (ST_DWithin, ST_KDE)', 'Open-source, reproducible alternative to ArcGIS; native support for large-scale spatial joins'],
        ['GeoPandas',   'Python-side spatial data manipulation and PostGIS interface',                       'Standard geospatial Python library; integrates with Census API and ACS data pipelines'],
        ['PyMC',        'Bayesian hierarchical spatial regression for coverage gap modeling',                 'Provides posterior uncertainty quantification unavailable in classical spatial lag regression'],
        ['Python 3.11', 'Primary implementation language for all pipeline components',                       'Cross-platform; extensive geospatial and ML ecosystem'],
        ['HPC Cluster', 'Batch execution of approx. 3,000 API calls and GPT inference at scale',            'CU Boulder Research Computing; enables parallel ZIP-code scanning and reduces wall-clock time'],
    ],
    widths=[1.2, 2.9, 2.1]
)

H('4.3  Multi-Agent System (MAS) Architecture', size=12)
B(
    'The data collection and classification pipeline is implemented as a '
    'Multi-Agent System (MAS) orchestrated via LangChain. Rather than a single '
    'monolithic script, the system decomposes the workflow into four specialized '
    'agents, each with a defined input contract, output schema, and failure-handling '
    'protocol. LangChain manages agent instantiation, inter-agent message passing, '
    'shared memory (deduplication state), and tool routing. This architecture '
    'provides three practical advantages: (1) each agent can be tested, modified, '
    'or replaced independently without disrupting the full pipeline; (2) agents '
    'can be parallelized at the ZIP-code level on HPC infrastructure; and '
    '(3) the modular design makes the framework portable to other cities or '
    'business categories by substituting only the configuration layer.'
)
B('The four agents and their responsibilities are as follows:')
table(
    headers=['Agent', 'Responsibility', 'Primary Tool', 'Output'],
    rows=[
        ['Search Agent',
         'Query Google Maps Places API by ZIP code; iterate paginated results; deduplicate records by place_id across all ZIP codes',
         'Google Maps Places API',
         'Deduplicated list of place_id references with source ZIP metadata'],
        ['Enrichment Agent',
         'Fetch Place Details per unique place_id (address, coordinates, hours, reviews, attributes); cross-validate open/closed status via Yelp Fusion API',
         'Google Maps Place Details API; Yelp Fusion API',
         'Enriched record dict with 15 raw attribute fields per establishment'],
        ['Classification Agent',
         'Submit enriched record to GPT-4o-mini with structured NAICS prompt (temp. = 0.0); parse JSON response; flag low-confidence assignments',
         'OpenAI GPT-4o-mini',
         'Classified record with Calculated_NAICS, Confidence, Reasoning, Employees_Estimated, Year_Established'],
        ['QA Agent',
         'Validate output schema completeness; detect duplicate names within 200-meter radius; flag records where AI confidence is Low and Yelp disagrees on open status; route flagged records for manual review',
         'PostGIS ST_DWithin; rule-based checks',
         'Validated, quality-scored dataset exported to PostGIS and timestamped CSV (22 columns, NETS-aligned)'],
    ],
    widths=[1.3, 2.6, 1.5, 1.8]
)
B(
    'The Classification Agent prompt specifies NAICS code 446110, provides '
    'distinguishing criteria separating pharmacies from health food stores, medical '
    'clinics, and vitamin shops, and requires a JSON-structured response. '
    'Temperature is set to 0.0 across all LLM calls, ensuring deterministic and '
    'reproducible outputs. The complete MAS source code is openly available at '
    'github.com/McDonaldCrispyThigh/NETS-AI.'
)

H('4.4  Ground Truth Construction', size=12)
B(
    'The Minnesota Board of Pharmacy licensure database will be obtained via formal '
    'data request to mn.gov/boards/pharmacy/public/datarequests.jsp. Records will be '
    'filtered to active retail pharmacy licenses in Minneapolis ZIP codes and geocoded '
    'using the U.S. Census Bureau Geocoder API. This constitutes the authoritative '
    'reference population. NCPDP DataQ records serve as a supplementary verification '
    'layer. Chain pharmacies are identified by name matching against a predefined list '
    '(CVS, Walgreens, Rite Aid, Target Pharmacy, Walmart Pharmacy, Costco Pharmacy); '
    'remaining establishments are classified as independent.'
)

H('4.5  Validation Framework', size=12)
B('Three pairwise comparisons are conducted:')
bullet('Comparison A (AI vs. Ground Truth):',
       'AI records spatially matched to ground-truth pharmacies within 100 meters. '
       'Metrics: Coverage Rate (Recall), False Discovery Rate, F1 score. '
       'Subgroup analysis: chain vs. independent; high-income vs. low-income ZIP codes.')
bullet('Comparison B (NETS vs. Ground Truth):',
       'Same matching protocol applied to NETS records active in 2024 to 2025. '
       'Produces a direct benchmark for NETS recall and precision comparable to the '
       'AI agent.')
bullet('Comparison C (AI vs. NETS):',
       'Establishments present in AI data but absent from NETS (AI-only) and vice '
       'versa (NETS-only) are characterized by type, location, and socioeconomic '
       'context of the surrounding census tract.')

H('4.6  Spatial Analysis with PostGIS and GeoPandas', size=12)
B(
    'Spatial data management and querying are implemented in PostGIS, a spatial '
    'extension for PostgreSQL. All three datasets (AI-generated, NETS, and ground '
    'truth) are ingested into a PostGIS database projected in NAD83 / UTM Zone 15N. '
    'Spatial matching operations use ST_DWithin for the 100-meter proximity threshold; '
    'kernel density estimation is computed via ST_KernelDensity and visualized '
    'through GeoPandas. PostGIS enables reproducible, version-controlled spatial '
    'queries that can be re-executed without proprietary software licenses, '
    'consistent with the open-science orientation of this thesis.'
)
B(
    'Analysis includes: kernel density maps of pharmacy locations per dataset, '
    'overlaid to identify spatial clusters of divergence; spatial join of coverage '
    'gap indicators with ACS 2019-2023 tract-level data (median household income, '
    'percent non-white, percent uninsured); and bivariate choropleth mapping of '
    'ground-truth coverage against NETS and AI coverage gaps at the census tract level.'
)

H('4.7  Bayesian Spatial Regression with PyMC', size=12)
B(
    'To model the probability of undercounting as a function of neighborhood '
    'characteristics, this study employs a Bayesian hierarchical spatial regression '
    'implemented in PyMC. The outcome variable is a binary indicator per census '
    'tract: whether AI or NETS undercounts relative to ground truth. Predictors '
    'include median household income, percent non-white population, percent '
    'uninsured, and pharmacy density (pharmacies per square kilometer).'
)
B(
    'The Bayesian framework offers two advantages over classical spatial lag '
    'regression. First, it produces full posterior distributions over regression '
    'coefficients, quantifying uncertainty in the relationship between neighborhood '
    'disadvantage and undercounting probability rather than reporting only point '
    'estimates. Second, a Conditional Autoregressive (CAR) prior is specified for '
    'the spatial random effect, directly modeling spatial autocorrelation in '
    'undercounting patterns across adjacent census tracts. Model convergence is '
    'assessed via R-hat statistics and posterior predictive checks. Results are '
    'reported as posterior mean coefficients with 95 percent highest density '
    'intervals.'
)

H('4.8  Computational Infrastructure', size=12)
B(
    'The full MAS pipeline across approximately 40 ZIP codes generates an estimated '
    '3,000 to 5,000 Google Maps API calls (search plus detail), up to 3,500 GPT '
    'inference calls, and associated Yelp cross-validation requests. To manage '
    'throughput and rate limits efficiently, the pipeline is deployed on the '
    'University of Colorado Boulder Research Computing HPC cluster. ZIP-code-level '
    'Search and Enrichment agents are parallelized across compute nodes; '
    'Classification Agent calls are batched to respect OpenAI rate limits. '
    'PostGIS runs on a persistent database instance on the same cluster, eliminating '
    'file-transfer overhead between pipeline stages. All HPC jobs are submitted '
    'via SLURM with deterministic random seeds and logged API call records, '
    'ensuring full reproducibility of the collection run.'
)

H('4.9  Interpretive Framework', size=12)
B(
    'Coverage gap patterns across the three datasets are interpreted through the '
    'following analytical matrix. Divergence is treated not as noise to minimize '
    'but as a signal to decode: the spatial pattern of each dataset\'s errors '
    'reveals the structural biases of its underlying data collection mechanism.'
)
B('Summary of divergence interpretations:')
table(
    headers=['AI Coverage', 'NETS Coverage', 'Interpretation'],
    rows=[
        ['Below ground truth', 'Below ground truth',  'Shared structural gap: both miss very small or very new establishments'],
        ['Near ground truth',  'Below ground truth',  'NETS-specific bias: D and B credit-file gap for independent pharmacies'],
        ['Below ground truth', 'Near ground truth',   'AI-specific gap: Google Maps listing absence or stale active listing'],
        ['Gaps in low-income areas', 'Gaps in low-income areas', 'Systematic equity bias present in both datasets'],
    ],
    widths=[1.5, 1.5, 3.2]
)

# ════════════════════════════════════════════════════════════════
# 5. CONTRIBUTIONS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
H('5.  Expected Contributions')

H('5.1  Empirical Contribution', size=12)
B(
    'The first systematic comparison of AI-generated pharmacy establishment data '
    'against both regulatory ground truth and NETS in a U.S. city. This provides '
    'concrete, quantified estimates of NETS data quality for a health-critical '
    'business category, filling a two-decade gap in the pharmacy desert literature.'
)

H('5.2  Methodological Contribution', size=12)
B(
    'This thesis delivers a portable audit pipeline: an open-source, modular '
    'Multi-Agent System that any researcher can deploy against any U.S. city or '
    'commercial establishment category to independently audit the accuracy of '
    'proprietary establishment databases. The pipeline requires only API credentials '
    'and a category configuration file to transfer to a new context; the MAS '
    'architecture, LangChain orchestration, PostGIS spatial backend, and PyMC '
    'inference model are fully reusable without modification.'
)
B(
    'Marginal cost per additional city is below USD 100, compared to the four-figure '
    'institutional licensing fees required for NETS access. This cost structure '
    'democratizes database auditing: a capacity previously available only to '
    'well-funded research teams with proprietary data agreements becomes accessible '
    'to any university researcher or public health practitioner. The pipeline is '
    'documented at github.com/McDonaldCrispyThigh/NETS-AI and is designed for '
    'replication, extension, and peer review, directly addressing the reproducibility '
    'deficit in commercial geography research.'
)

H('5.3  Policy Contribution', size=12)
B(
    'If NETS systematically undercounts independent pharmacies in lower-income '
    'Minneapolis neighborhoods, existing pharmacy desert analyses may have '
    'understated the severity of the access crisis. This finding argues for '
    'methodological revision in pharmacy access research and for policy '
    'interventions calibrated to a more accurate picture of geographic disadvantage.'
)

H('5.4  Conceptual Contribution', size=12)
B(
    'This study advances the concept of AI-generated data as a methodological probe: '
    'a reproducible instrument for exposing the assumptions and biases embedded in '
    'existing proprietary data infrastructure. This framing has broad applicability '
    'across commercial geography, public health informatics, and urban data science.'
)

# ════════════════════════════════════════════════════════════════
# 6. LIMITATIONS
# ════════════════════════════════════════════════════════════════
H('6.  Limitations and Mitigation Strategies')
table(
    headers=['Limitation', 'Nature', 'Mitigation Strategy'],
    rows=[
        ['Cross-sectional scope',
         'AI pipeline captures point-in-time data only; NETS longitudinal depth cannot be replicated',
         'Focus on 2024 to 2025 cross-section; use NETS time-series as descriptive context'],
        ['Estimated metadata',
         'Year_Established and Employees_Estimated are GPT inferences without documentary verification',
         'Label as estimates; exclude from primary validation metrics; sensitivity analysis'],
        ['Google Maps bias',
         'User-contributed listings may under-represent businesses with low digital presence',
         'Cross-validate with Yelp API; flag establishments absent from both sources'],
        ['Temporal mismatch',
         'Ground truth, NETS, and AI data may not share identical reference dates',
         'Apply active/inactive filters consistently; document reference date per source'],
        ['Single-MSA scope',
         'Findings are bounded to the Minneapolis-St. Paul metropolitan area and may not generalize to smaller cities, rural areas, or metros with different pharmacy market structures',
         'Discuss boundary conditions explicitly; design pipeline for direct replication in comparable MSAs (e.g., Chicago, Detroit) in future work'],
    ],
    widths=[1.4, 2.3, 2.5]
)

# ════════════════════════════════════════════════════════════════
# 7. TIMELINE
# ════════════════════════════════════════════════════════════════
H('7.  Research Timeline')
table(
    headers=['Phase', 'Tasks', 'Timeline'],
    rows=[
        ['1: Setup',
         'Add pharmacy config to agent; submit MN Board of Pharmacy data request; request NETS extract via U of Minnesota; confirm NCPDP DataQ access',
         'April 2026'],
        ['2: Collection',
         'Run MAS pipeline across approx. 40 Minneapolis-St. Paul MSA ZIP codes on HPC; collect and deduplicate full pharmacy dataset',
         'May 2026'],
        ['3: Ground Truth',
         'Geocode Board of Pharmacy records; verify against NCPDP DataQ; establish spatial matching protocol',
         'May 2026'],
        ['4: Comparisons',
         'Execute Comparisons A, B, and C; compute precision, recall, and F1 per subgroup',
         'June 2026'],
        ['5: Spatial Analysis',
         'PostGIS kernel density and spatial joins; PyMC Bayesian hierarchical regression; choropleth mapping via GeoPandas',
         'June to July 2026'],
        ['6: Writing',
         'Draft Chapters 1 through 5; incorporate advisor feedback; revise and refine',
         'July to September 2026'],
        ['7: Submission',
         'Final revision; honors thesis submission; prepare conference abstract for submission',
         'October 2026'],
    ],
    widths=[1.3, 3.7, 1.2]
)

# ════════════════════════════════════════════════════════════════
# 8. BUDGET
# ════════════════════════════════════════════════════════════════
H('8.  Budget Estimate')
table(
    headers=['Item', 'Estimated Cost (USD)'],
    rows=[
        ['Google Maps Places API (approx. 3,500 records across MSA, search plus detail calls)', 'approx. 130'],
        ['OpenAI GPT-4o-mini (approx. 3,500 classifications at 0.0002 per call)',              'approx. 0.70'],
        ['Yelp Fusion API',                                                                     'Free (basic tier)'],
        ['PostGIS / PostgreSQL',                                                                'Free, open-source'],
        ['PyMC',                                                                                'Free, open-source'],
        ['HPC compute time (CU Boulder Research Computing)',                                    'Free (institutional allocation)'],
        ['MN Board of Pharmacy and Ramsey County data request fees',                            'approx. 50'],
        ['Total',                                                                               'approx. 180'],
    ],
    widths=[4.5, 1.7]
)
B(
    'The low marginal cost of this portable audit pipeline, relative to NETS '
    'four-figure institutional licensing fees, is itself a statement about data '
    'accessibility in research infrastructure. The PostGIS, PyMC, and LangChain '
    'components are fully open-source; HPC access is available through standard '
    'university research computing allocations at no direct cost. A replication '
    'for any additional U.S. city or business category could be conducted for '
    'under USD 200, enabling systematic national-scale auditing of commercial '
    'database quality at a fraction of the cost of current approaches.'
)

# ════════════════════════════════════════════════════════════════
# 9. REFERENCES
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
H('9.  References')

refs = [
    'Barnatchez, K., Crane, L. D., and Decker, R. (2017). An Assessment of the National Establishment Time Series (NETS) Database. Finance and Economics Discussion Series 2017-110. Federal Reserve Board, Washington, D.C. Available at federalreserve.gov.',
    'Chapple, K. and Jacobus, R. (2009). Retail trade as a route to neighborhood revitalization. In N. Pindus, H. Wial, and H. Wolman (Eds.), Urban and Regional Policy and Its Effects. Brookings Institution Press.',
    'Dill, M. J. and Gelmon, S. (2023). Geographic access to pharmacies in the United States. Journal of the American Pharmacists Association, 63(1), 148-156.',
    'Federal Reserve Bank of Minneapolis. (2017). An Assessment of the National Establishment Time Series Database. Working Paper WP17-29. minneapolisfed.org.',
    'Federal Reserve Board. (2019). Finance and Economics Discussion Series 2019-034. Washington, D.C.',
    'Fox9 Minneapolis. (2024, February 13). MN pharmacy closures accelerating, report warns of pharmacy desert increase. fox9.com.',
    'Meltzer, R. and Schuetz, J. (2012). Bodegas or bagel shops? Neighborhood differences in retail and household services. Economic Development Quarterly, 26(1), 73-94.',
    'Minnesota Department of Health. (2024). Pharmacy Deserts in Minnesota 2009-2024. St. Paul, MN: Minnesota Department of Health.',
    'MinnPost. (2025, March). There\'s a Pharmacy Shortage in Minnesota. minnpost.com.',
    'National Community Pharmacists Association (NCPA). (2026). State of the Indy Report. ncpanet.org.',
    'National Council for Prescription Drug Programs (NCPDP). (2024). NCPDP DataQ: Next-Generation Pharmacy Database. ncpdp.org.',
    'Neumark, D., Zhang, J., and Wall, B. (2011). Employment dynamics and business relocation: New evidence from the National Establishment Time-Series. Research in Labor Economics, 32, 39-83.',
    'Pereira, C. et al. (2024). The Minnesota Pharmacy Landscape Is Drying Up: Pharmacy Desert Emergence and Implications. Abstract 551580. American Public Health Association Annual Meeting, Minneapolis.',
    'PLOS ONE. (2025). Distribution of pharmacy deserts and its association with digital divide and residential redlining across the United States. doi:10.1371/journal.pone.0330027.',
    'Qato, D. M., Daviglus, M. L., Wilder, J., Lee, T., Qato, D., and Lambert, B. (2014). Pharmacy deserts are prevalent in Chicago\'s predominantly minority communities, raising medication access concerns. Health Affairs, 33(8), 1359-1367.',
    'Shen, Y. et al. (2025). On the Use of LLMs for GIS-Based Spatial Analysis. ISPRS International Journal of Geo-Information, 14(10), 401.',
    'ScienceDirect. (2025). Predicting restaurant survival using nationwide Google Maps data. Available at sciencedirect.com.',
    'ScienceDirect. (2025). Large Language Models in Urban Planning: A Systematic Review and Conceptual Framework. Journal of Urban Technology. doi:10.1080/10630732.2025.2556551.',
    'Star Tribune. (2024). Number of pharmacy deserts grow as chain stores close. startribune.com.',
    'University of Minnesota College of Pharmacy. (2024). Mapping pharmacy deserts across Minnesota. pharmacy.umn.edu.',
    'Walls, D. W. (2007). National Establishment Time-Series Database: Data Overview. SSRN Working Paper 1022962.',
    'Washington State Department of Commerce. (2021). Data Driven Insight: Evaluating the Dun and Bradstreet Toolkit. Olympia, WA.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Inches(0.4)
    p.paragraph_format.first_line_indent  = Inches(-0.4)
    p.paragraph_format.space_after        = Pt(4)
    r = p.add_run(ref)
    font(r, size=10)

# ── Save ────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Research_Proposal.docx')
doc.save(out)
print('Saved:', out)
