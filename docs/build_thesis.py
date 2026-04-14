"""
build_thesis.py
Generate docs/Honors_Thesis_Zheng_2026.docx
Run: python docs/build_thesis.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.different_first_page_header_footer = True

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

sectPr = section._sectPr
pgNumType = OxmlElement('w:pgNumType')
pgNumType.set(qn('w:start'), '0')
sectPr.append(pgNumType)

# ── Helpers ─────────────────────────────────────────────────────────────────

def add_page_number_to_footer(sec):
    footer = sec.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = para.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fc1)
    run2 = para.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._r.append(instr)
    run3 = para.add_run()
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fc2)


def _run(para, text, bold=False, italic=False, size=12):
    r = para.add_run(text)
    r.font.name   = 'Times New Roman'
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    return r


def chapter_heading(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(24)
    pf.space_after       = Pt(12)
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _run(p, text, bold=True, size=14)
    return p


def section_heading(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(12)
    pf.space_after       = Pt(6)
    pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _run(p, text, bold=True, size=12)
    return p


def body(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Inches(0.5)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _run(p, text)
    return p


def ref_entry(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.left_indent       = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _run(p, text)
    return p


def title_center(text, size=12, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    _run(p, text, bold=bold, italic=italic, size=size)
    return p


add_page_number_to_footer(section)


def add_figure(img_path, label, caption_text):
    """Insert a centered figure with a caption line below it."""
    fig_para = doc.add_paragraph()
    pf = fig_para.paragraph_format
    pf.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(18)
    pf.space_after  = Pt(4)
    run = fig_para.add_run()
    full_path = os.path.join(os.path.dirname(__file__), '..', img_path)
    run.add_picture(full_path, width=Inches(5.5))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(18)
    r1 = cap.add_run(label + '.  ')
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    r1.font.bold = True
    r2 = cap.add_run(caption_text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

title_center(
    'Can Artificial Intelligence Reconstruct Commercial Geographies?',
    size=16, bold=True, space_after=4
)
title_center(
    'Evaluating AI-Generated Pharmacy Establishment Data as a Validator of the '
    'National Establishment Time-Series (NETS) Database in Minneapolis-St. Paul',
    size=13, italic=True, space_after=16
)
title_center(
    'An Honors Thesis submitted in partial fulfillment of the requirements for the '
    'degree of Bachelor of Arts in Geography and Bachelor of Science in '
    'Applied Mathematics',
    size=12, space_after=20
)

for lbl, val in [
    ('Author:', 'Congyuan Zheng'),
    ('Department:', 'Department of Geography, University of Colorado Boulder'),
    ('Primary Advisor:', 'Prof. Jessica Finlay, University of Colorado Boulder'),
    ('Co-Advisor:', 'Prof. Michael H. Esposito, University of Minnesota'),
    ('Date:', 'April 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    _run(p, lbl + '  ', bold=True)
    _run(p, val)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Acknowledgments')

body(
    'This thesis would not have been possible without the guidance, generosity, and '
    'intellectual rigor of my advisory team. Prof. Jessica Finlay provided the '
    'foundational framing for this project, consistently pushing the research toward '
    'its most meaningful contribution at the intersection of urban geography and health '
    'equity. Her mentorship shaped not only the thesis question but the approach to '
    'evidence that runs throughout this document.'
)
body(
    'Prof. Michael H. Esposito brought an essential empirical grounding to the project, '
    'challenging early assumptions about NETS data quality and directing my attention '
    'toward the Federal Reserve Board assessment literature that became a cornerstone '
    'of the theoretical framework. His expertise in the sociology of racial health '
    'disparities and Bayesian demographic methods anchored the methodological design '
    'in existing scholarship and prevented the research from treating the AI pipeline '
    'as an end rather than a means.'
)
body(
    'Yue Sun provided invaluable technical mentorship throughout the implementation '
    'phase, offering guidance on spatial analysis methods, GeoPandas workflows, and '
    'the interpretation of fuzzy matching results. Conversations with Yue were '
    'consistently clarifying and productive.'
)
body(
    'I am grateful to the University of Colorado Boulder Department of Geography for '
    'supporting this work through the Honors Program, and to the CU Boulder Library '
    'for providing access to the data and literature resources on which this research '
    'depends. All errors and limitations in this document are my own.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (text)
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Table of Contents')

for entry, pg in [
    ('Acknowledgments', 'ii'),
    ('Abstract', 'iii'),
    ('Chapter 1: Introduction', '1'),
    ('  1.1  The Pharmacy Desert Crisis in Minnesota', '1'),
    ('  1.2  The Problem: NETS as a Black Box', '3'),
    ('  1.3  Research Opportunity and Contributions', '5'),
    ('Chapter 2: Literature Review', '7'),
    ('  2.1  Pharmacy Access and Health Equity', '7'),
    ('  2.2  NETS Database: Strengths and Documented Limitations', '9'),
    ('  2.3  AI Agents and LLMs in Urban Geography', '11'),
    ('  2.4  Ground Truth and Regulatory Data Sources', '13'),
    ('Chapter 3: Methodology', '15'),
    ('  3.1  Study Area', '15'),
    ('  3.2  Multi-Agent System Architecture', '16'),
    ('  3.3  Ground Truth Construction', '18'),
    ('  3.4  Validation Framework', '20'),
    ('  3.5  Spatial Analysis', '21'),
    ('  3.6  Limitations', '22'),
    ('Chapter 4: Results', '23'),
    ('  4.1  AI Data Collection Results', '23'),
    ('  4.2  Validation Against NPPES', '25'),
    ('  4.3  Pharmacy Desert Spatial Analysis', '28'),
    ('  4.4  North Minneapolis Case Study', '30'),
    ('Chapter 5: Discussion', '32'),
    ('  5.1  What Data Source Discrepancies Reveal', '32'),
    ('  5.2  Implications for Pharmacy Desert Research', '34'),
    ('  5.3  The Portable Audit Pipeline as a Contribution', '36'),
    ('Chapter 6: Conclusion', '38'),
    ('References', '41'),
    ('List of Figures', '47'),
]:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r1 = p.add_run(entry)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run('\t' + pg)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Abstract')

body(
    'The National Establishment Time-Series (NETS) database serves as a primary data '
    'infrastructure for urban commercial geography research, yet its construction from '
    'proprietary Dun and Bradstreet credit records introduces persistent coverage gaps '
    'that receive limited critical examination in the health geography literature. A '
    '2017 Federal Reserve Board assessment confirmed that NETS exhibits its largest '
    'discrepancies among small establishments, operates on a two-to-three-year lag in '
    'recording business deaths, and contains documented underrepresentation of '
    'minority-owned enterprises (Barnatchez, Crane, and Decker, 2017).'
)
body(
    'The timing matters. Pharmacy desert research is accelerating precisely as its '
    'primary data sources are most likely to mislead. Eight percent of Minnesota '
    'residents now live in a pharmacy desert, up from 6.2 percent in 2009. In North '
    'Minneapolis, Walgreens closed its West Broadway Avenue location in February 2023. '
    'CVS followed with four additional metropolitan area closures. Approximately '
    '34 percent of Twin Cities residents live in or near a pharmacy desert '
    '(Minnesota Department of Health, 2024). When the databases used to track these '
    'closures lag two to three years behind reality, policy responses lag with them.'
)
body(
    'This thesis develops and deploys a portable, reproducible Multi-Agent System that '
    'collects, classifies, and structures pharmacy establishment data for the '
    'Minneapolis-St. Paul Metropolitan Statistical Area using the Google Maps Places '
    'API and OpenAI GPT-4o-mini, then validates the resulting dataset against the '
    'NPPES NPI Registry as an interim ground truth. The pipeline recovered 399 unique '
    'pharmacy records across 101 ZIP codes, achieving a 99.2 percent NAICS 446110 '
    'match rate at 94.5 percent high-confidence classification.'
)
body(
    'Validation against NPPES revealed 717 false negatives, of which 331 (46.2 '
    'percent) were corporate legal name entries rather than individual retail '
    'locations, and 95 (13.2 percent) were records for closed or acquired chains still '
    'marked active in the registry. Spatial analysis identified pharmacy deserts in '
    '83.9 percent of 472 census tracts analyzed, with North Minneapolis exhibiting a '
    'desert rate of 88.9 percent across 18 tracts. The pipeline operates at a marginal '
    'cost below USD 100 per metropolitan area and is publicly available at '
    'github.com/McDonaldCrispyThigh/NETS-AI.'
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
_run(p, 'Keywords: ', bold=True)
_run(p, 'NETS Database | Pharmacy Deserts | Health Equity | AI Agents | Minneapolis | '
    'Reproducibility | Urban Commercial Geography', italic=True)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Chapter 1: Introduction')

body(
    'Urban pharmacy access sits at the intersection of commercial geography, public '
    'health infrastructure, and spatial equity. The locations of retail pharmacies '
    'determine whether residents can obtain prescription medications within a walkable '
    'distance, access immunization services, or receive over-the-counter health '
    'counseling from a trained pharmacist. These functions are most critical for '
    'populations with limited mobility, vehicle access, or primary care contact, '
    'populations that are disproportionately low-income, elderly, and racial minorities. '
    'When pharmacies close or never open in these communities, the resulting gaps '
    'constitute a structural health equity problem with measurable consequences for '
    'chronic disease management and preventive care.'
)
body(
    'This thesis examines two interconnected problems. The first is the growing '
    'pharmacy desert crisis in the Minneapolis-St. Paul metropolitan area, where '
    'chain pharmacy retreat and independent pharmacy closure have left significant '
    'portions of the urban core without accessible retail dispensing services. The '
    'second is the measurement problem that underlies pharmacy desert research: the '
    'commercial establishment databases used to identify and track pharmacy access '
    'contain structural errors that have not been critically examined in the health '
    'geography literature, and those errors are directionally concentrated in the '
    'communities most affected by the access crisis. Addressing both problems requires '
    'an independent data source, and this thesis develops that source.'
)

section_heading('1.1 The Pharmacy Desert Crisis in Minnesota')

body(
    'Access to community pharmacies represents one of the most fundamental determinants '
    'of medication adherence and population health outcomes. Pharmacies serve as '
    'first-contact healthcare providers, dispensing prescription medications, '
    'providing immunizations, and delivering over-the-counter health counseling in '
    'settings that reach populations who lack consistent access to primary care. The '
    'geographic distribution of pharmacies is therefore not a purely commercial '
    'question but a public health infrastructure concern with direct implications for '
    'health equity, particularly in low-income urban neighborhoods where chain '
    'pharmacies have retreated and independent operators have struggled to remain '
    'financially viable.'
)
body(
    'Minnesota is confronting an accelerating pharmacy desert crisis that has worsened '
    'measurably over the past fifteen years. A 2024 report by the Minnesota Department '
    'of Health documented that 8 percent of Minnesota residents now live in a pharmacy '
    'desert, defined as a census tract where the nearest retail pharmacy is more than '
    'half a mile from the tract centroid, an increase from 6.2 percent in 2009 '
    '(Minnesota Department of Health, 2024). Statewide, approximately 44 percent of '
    'all Minnesota pharmacies have closed over the past decade, a rate that '
    'substantially outpaces new openings. Approximately 34 percent of Twin Cities '
    'metropolitan area residents live in or near a pharmacy desert, a proportion that '
    'has grown as closures have concentrated in communities with the least commercial '
    'attractiveness for new entrants.'
)
body(
    'The distribution of closures is far from random. Independent pharmacies have '
    'accounted for approximately 60 percent of pharmacy closures over the past decade, '
    'despite disproportionately serving low-income, rural, and minority communities '
    'where chain pharmacies have limited commercial incentive to operate (National '
    'Community Pharmacists Association, 2026). The geographic pattern of closure '
    'concentrates in two types of communities: rural areas where population density '
    'cannot support any pharmacy business model, and urban neighborhoods of color '
    'where chains have retreated on commercial grounds. Both types of communities are '
    'disproportionately represented among the census tracts that qualify as pharmacy '
    'deserts under the Qato et al. (2014) half-mile access threshold.'
)
body(
    'The scale of the deficit in North Minneapolis is stark. Sixteen of 18 census '
    'tracts in ZIP codes 55411 and 55412 qualify as pharmacy deserts under the '
    'half-mile threshold. One retail pharmacy serves both ZIP codes.'
)
body(
    'This is the outcome of a decade of chain retreat. The Walgreens at 627 West '
    'Broadway Avenue, which had served ZIP code 55411 for decades, closed in February '
    '2023. CVS subsequently closed four additional metropolitan area locations. The '
    'communities left with reduced pharmacy access are predominantly Black, with '
    'median household incomes substantially below county averages. This study\'s '
    'spatial analysis places the desert rate at 88.9 percent across the 18 tracts, '
    'the most acute concentration in the MSA dataset.'
)
body(
    'The crisis is compounded by a measurement problem that has received insufficient '
    'attention in the health geography literature. Researchers and policymakers who '
    'seek to document and address pharmacy access disparities rely on commercial '
    'establishment databases whose construction and quality have rarely been '
    'subjected to critical scrutiny in this domain. If the database used to identify '
    'pharmacy deserts is itself inaccurate, the deserts identified may reflect data '
    'quality artifacts as much as real-world commercial geography. This study treats '
    'that measurement problem as a primary research object rather than a background '
    'assumption, and it develops a methodology for independent validation of '
    'establishment data at metropolitan scale.'
)

section_heading('1.2 The Problem: NETS as a Black Box')

body(
    'The National Establishment Time-Series database, assembled by Walls and Associates '
    'from Dun and Bradstreet commercial credit records, has served as the primary '
    'longitudinal establishment dataset in urban commercial geography for over two '
    'decades. Its applications span neighborhood retail change analysis (Chapple and '
    'Jacobus, 2009), food access research, pharmacy desert mapping, and business cycle '
    'studies. NETS covers approximately 45 million U.S. establishment records with '
    'annual snapshots dating to 1990, making it the most temporally comprehensive '
    'commercial establishment database available to academic researchers (Walls, 2007). '
    'The database\'s longitudinal architecture enables studies of business births, '
    'deaths, and relocations that no alternative publicly available dataset can support '
    'at comparable scope and historical depth.'
)
body(
    'NETS is constructed from an underlying source designed for commercial credit '
    'assessment rather than geographic or population research. A rigorous 2017 '
    'assessment by Federal Reserve Board economists Barnatchez, Crane, and Decker '
    'identified three structural limitations with direct relevance to pharmacy desert '
    'research. NETS exhibits its largest coverage gaps among small establishments, '
    'the category most directly relevant to independent pharmacy survival. The database '
    'operates on a two-to-three-year lag in recording business deaths, meaning that a '
    'pharmacy that closes in a given year may remain listed as active in NETS for '
    'several subsequent vintages. The Dun and Bradstreet source contains documented '
    'underrepresentation of minority-owned enterprises, a bias that concentrates '
    'coverage errors in the neighborhoods where pharmacy desert research is most '
    'consequential (Barnatchez, Crane, and Decker, 2017).'
)
body(
    'A subsequent evaluation by the Washington State Department of Commerce confirmed '
    'the minority-business underrepresentation pattern in a different geographic '
    'context, indicating that the limitation is not regionally idiosyncratic '
    '(Washington State Department of Commerce, 2021). These findings, while published '
    'in the economics literature, have not been consistently incorporated into health '
    'geography research that uses NETS as a primary data source. The gap between what '
    'NETS\'s known limitations predict and what pharmacy desert literature assumes '
    'about NETS data quality is the central intellectual problem this thesis addresses.'
)
body(
    'Both error types point in the same direction. NETS overstates pharmacy coverage '
    'in the communities experiencing the fastest chain retreat, and understates '
    'coverage among the independent operators most likely to serve communities of '
    'color. The result is a database that directionally minimizes the pharmacy desert '
    'problem in precisely the neighborhoods where that minimization carries the '
    'greatest policy consequences.'
)
body(
    'A researcher using NETS to identify pharmacy deserts in North Minneapolis in '
    '2024 or 2025 would likely find the closed Walgreens at 627 West Broadway still '
    'listed as active. The database says pharmacy coverage exists. The neighborhood '
    'says otherwise.'
)

section_heading('1.3 Research Opportunity and Contributions')

body(
    'The convergence of large language models, public business data infrastructure, '
    'and open spatial analysis tools creates a new methodological possibility for '
    'addressing the NETS transparency problem. Where commercial database construction '
    'relies on proprietary credit records and opaque pipelines, a researcher can now '
    'assemble an independent pharmacy census from public sources at a fraction of the '
    'cost of institutional database licensing. This study develops exactly that '
    'capability: a portable Multi-Agent System that queries the Google Maps Places API '
    'for all pharmacy locations within a metropolitan area, enriches each record with '
    'structured metadata, and submits the enriched records to GPT-4o-mini for NAICS '
    'classification, producing a dataset structurally comparable to NETS at a marginal '
    'cost below USD 100 per metropolitan area.'
)
body(
    'This research addresses four primary questions. The first question concerns '
    'coverage: to what extent does the AI-generated dataset capture the complete '
    'population of active Minneapolis-St. Paul pharmacies relative to regulatory '
    'ground truth? The second concerns accuracy: how accurately does the AI agent '
    'classify pharmacies by NAICS code 446110, and what establishment types are most '
    'frequently misclassified? The third concerns divergence: where the AI dataset '
    'and NETS diverge from regulatory ground truth, do their errors follow distinct '
    'spatial and organizational patterns? The fourth concerns equity implications: '
    'are identified biases spatially correlated with neighborhood-level socioeconomic '
    'disadvantage and racial composition?'
)
body(
    'The primary hypothesis is that NETS will exhibit lower recall relative to '
    'Minnesota Board of Pharmacy ground truth than the AI agent, due to the lag and '
    'coverage biases documented by Barnatchez, Crane, and Decker (2017), and that '
    'both datasets will exhibit their largest gaps in census tracts with below-median '
    'household income and above-median non-White population. The secondary hypothesis '
    'is that AI errors will be directionally opposite to NETS errors: the AI pipeline '
    'will miss low-digital-visibility independent pharmacies that NETS includes, '
    'while NETS will retain closed establishments that the AI pipeline correctly '
    'excludes as inactive.'
)
body(
    'The thesis contributes three distinct advances to urban geography scholarship. '
    'The empirical contribution is the first direct comparison of AI-generated '
    'pharmacy establishment data against regulatory ground truth across a full '
    'metropolitan statistical area, including a 717-record false negative '
    'decomposition that distinguishes corporate legal name artifacts from genuine '
    'missed retail locations. The methodological contribution is an open-source, '
    'modular pipeline deployable for any city and any business category at minimal '
    'cost. The policy contribution is a spatial analysis framework that connects '
    'data infrastructure quality to health equity outcomes, demonstrating that the '
    'measurement tools used in pharmacy desert research can themselves obscure the '
    'severity of the crisis they are designed to document.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Chapter 2: Literature Review')

body(
    'This chapter reviews four bodies of scholarship that inform the research design '
    'and interpretive framework of this thesis. The first concerns the academic '
    'literature on pharmacy access and health equity, which establishes the public '
    'health significance of the pharmacy desert problem and the spatial methodologies '
    'used to measure it. The second examines the documented strengths and limitations '
    'of the NETS database, drawing primarily on the Federal Reserve Board assessment '
    'and related commercial data quality studies. The third surveys the emerging '
    'literature on AI agents and large language models in urban geography, establishing '
    'the methodological feasibility of LLM-based establishment classification. The '
    'fourth reviews the landscape of ground truth and regulatory data sources '
    'available for pharmacy validation research.'
)

section_heading('2.1 Pharmacy Access and Health Equity')

body(
    'The scholarly foundation for pharmacy desert research was established by Qato '
    'et al. (2014), whose analysis of Chicago pharmacies introduced the half-mile '
    'buffer criterion that has since become the standard threshold in the literature. '
    'Qato and colleagues found that pharmacy deserts were concentrated in predominantly '
    'minority communities and that residents of these areas were significantly less '
    'likely to adhere to medication regimens for chronic conditions including '
    'hypertension and diabetes. The study\'s methodological contribution, a spatially '
    'explicit definition grounded in pedestrian access rather than administrative '
    'proximity, provided a reproducible framework that subsequent researchers have '
    'applied across dozens of metropolitan areas.'
)
body(
    'The half-mile threshold reflects the maximum walking distance most urban '
    'residents can reasonably sustain for routine medication pickup, a parameter '
    'validated by transportation research on pedestrian behavior in low-vehicle-access '
    'neighborhoods. Subsequent studies have confirmed the threshold\'s ecological '
    'validity: neighborhoods classified as pharmacy deserts under the half-mile '
    'standard show measurably lower rates of prescription fill completion and higher '
    'rates of medication nonadherence even after controlling for income, insurance '
    'status, and disease severity. These dose-response relationships strengthen the '
    'case that pharmacy desert classification captures a real access deficit rather '
    'than an arbitrary geographic artifact.'
)
body(
    'Minnesota-specific research has extended these findings with increasing '
    'granularity. The University of Minnesota College of Pharmacy tracked pharmacy '
    'desert prevalence across Minnesota from 2009 through 2024, documenting a steady '
    'increase confirmed in the Minnesota Department of Health\'s 2024 statewide '
    'report: 8 percent of state residents now live in a pharmacy desert, up from '
    '6.2 percent in 2009 (Minnesota Department of Health, 2024; Pereira et al., 2024). '
    'The longitudinal tracking makes Minnesota one of the few states with a '
    'longitudinal, multi-decade record of pharmacy desert dynamics, providing an '
    'unusually robust empirical baseline against which this study\'s AI-generated '
    'data can be evaluated.'
)
body(
    'A 2025 PLOS ONE study on pharmacy desert distribution linked geographic pharmacy '
    'absence to digital access divides and historical residential redlining patterns, '
    'establishing that current pharmacy deserts follow lines of historical '
    'disinvestment that systematically disadvantaged communities of color (PLOS ONE, '
    '2025). This finding connects the spatial patterns documented in Qato et al. '
    '(2014) to longer historical processes of neighborhood disinvestment, suggesting '
    'that pharmacy access disparities are structurally produced rather than '
    'commercially incidental. The historical redlining connection also raises the '
    'possibility that the digital visibility bias in AI-collected data may trace the '
    'same structural contours as the redlining footprint, because neighborhoods with '
    'lower internet adoption and fewer Google Business Profile claims may correlate '
    'with historically redlined areas.'
)
body(
    'The health consequences of pharmacy deserts extend beyond medication adherence. '
    'Pharmacies serve as first-line immunization providers, chronic disease management '
    'consultants, and emergency health referral points, particularly in communities '
    'where primary care access is limited. Dill and Gelmon (2023) examined geographic '
    'pharmacy access across the United States and found that the populations most '
    'likely to live in pharmacy deserts were disproportionately elderly, low-income, '
    'and racial minorities, precisely the populations with the highest chronic disease '
    'burden and the greatest dependence on regular pharmacy contact. The convergence '
    'of elevated disease burden, limited primary care access, and pharmacy desert '
    'geography in the same communities transforms what might initially appear as a '
    'commercial location problem into a structural health equity emergency requiring '
    'coordinated policy intervention.'
)

section_heading('2.2 NETS Database: Strengths and Documented Limitations')

body(
    'NETS has been the dominant data infrastructure for urban commercial geography '
    'research for more than two decades. Chapple and Jacobus (2009) demonstrated its '
    'analytical value in neighborhood retail change research, and Meltzer and Schuetz '
    '(2012) used NETS to document persistent differences in retail service distribution '
    'across income-stratified Chicago neighborhoods. The database\'s longitudinal '
    'architecture, covering over thirty years of annual snapshots with consistent '
    'establishment identifiers, enables the kind of business birth, death, and '
    'relocation analysis that no alternative publicly available dataset can support '
    'at comparable scope and historical depth. These strengths have made NETS the '
    'default choice in studies of urban commercial geography, food access, and '
    'pharmacy access.'
)
body(
    'The 2017 Federal Reserve Board assessment by Barnatchez, Crane, and Decker '
    'remains the most rigorous evaluation of NETS data quality available in the '
    'peer-reviewed literature. Their analysis compared NETS coverage against the '
    'Census Bureau\'s Longitudinal Business Database and identified three consistent '
    'patterns. Coverage of establishments with fewer than five employees was '
    'substantially lower in NETS than in the Census reference, with the gap largest '
    'in low-income neighborhoods and communities of color. Business deaths were '
    'recorded on a two-to-three-year lag, meaning establishments that ceased '
    'operations continued to appear in NETS as active for multiple subsequent '
    'vintages. And Dun and Bradstreet exhibited documented underrepresentation of '
    'minority-owned enterprises not explained by industry or size differences '
    '(Barnatchez, Crane, and Decker, 2017).'
)
body(
    'Neumark, Zhang, and Wall (2011) provided an earlier validation of NETS using '
    'state unemployment insurance records and found generally strong performance for '
    'large establishments while noting elevated error rates for the smallest employer '
    'size classes. Their findings prefigured the Barnatchez, Crane, and Decker '
    'assessment and established a pattern that multiple independent evaluations have '
    'confirmed: NETS performs best for the establishments that need data quality '
    'attention the least, the large chains with extensive credit records, and performs '
    'worst for the small, independent operators whose survival or closure is most '
    'consequential for community-level access analysis.'
)
body(
    'The Washington State Department of Commerce independently confirmed the '
    'minority-business underrepresentation finding in a 2021 evaluation of Dun and '
    'Bradstreet data applied to Washington state businesses, indicating that the '
    'pattern is not geographically idiosyncratic (Washington State Department of '
    'Commerce, 2021). This cross-state replication strengthens the inference that '
    'NETS coverage gaps in minority-owned small businesses are a structural feature '
    'of the underlying credit record system rather than a regional artifact, and that '
    'pharmacy desert research conducted with NETS data in any metropolitan area may '
    'carry the same undercount of independent operators serving communities of color.'
)
body(
    'These limitations have been largely absent from the health geography literature '
    'that uses NETS for pharmacy desert research. Studies that find fewer active '
    'pharmacies in North Minneapolis than in higher-income suburban neighborhoods may '
    'be observing real disparities, NETS lag effects, or some combination of both, '
    'with the relative contribution of each factor currently unknowable without an '
    'independent data source for comparison. The present study addresses this gap '
    'directly by constructing that independent source and conducting the comparison '
    'at full metropolitan scale.'
)

section_heading('2.3 AI Agents and LLMs in Urban Geography')

body(
    'Large language model applications in urban planning and geography have expanded '
    'rapidly since 2023. A 2025 systematic review in the Journal of Urban Technology '
    'identified over 80 peer-reviewed studies applying LLMs to urban planning tasks '
    'including land use classification, zoning analysis, and community feedback '
    'synthesis, finding competitive performance with domain expert annotations across '
    'a range of spatial classification tasks (ScienceDirect, 2025). The pace of '
    'methodological development in this area suggests that LLM-assisted spatial '
    'analysis will become a standard component of the urban geography researcher\'s '
    'toolkit within the next research generation, alongside existing tools such as '
    'GIS, spatial econometrics, and remote sensing.'
)
body(
    'Shen et al. (2025) specifically examined LLMs applied to GIS-based spatial '
    'analysis and found that GPT-4 class models achieved classification accuracy '
    'comparable to trained geographic information specialists when provided with '
    'structured prompts, geospatial context, and established taxonomic frameworks '
    'such as NAICS codes. Their study validated a core methodological assumption '
    'underlying this thesis: that LLMs can reliably assign industry classification '
    'codes to business establishments when given sufficient contextual information '
    'in a deterministic inference setting. The deterministic requirement is met in '
    'this study by setting GPT-4o-mini to temperature 0.0, which produces consistent '
    'outputs for identical inputs and is a necessary condition for scientific '
    'reproducibility.'
)
body(
    'The use of Google Maps as a research-grade data source has been validated in '
    'recent commercial geography literature. A 2025 ScienceDirect study predicting '
    'restaurant survival across a national sample used Google Maps listing data as '
    'the primary establishment inventory source, demonstrating that Google\'s business '
    'listings provide sufficiently complete coverage for urban commercial geography '
    'analysis in contexts where establishment counts exceed approximately twenty per '
    'ZIP code (ScienceDirect, 2025). This density threshold matters for the present '
    'study: pharmacies in the Minneapolis-St. Paul MSA average three to twelve '
    'establishments per ZIP code, below the density range where the API\'s 60-result '
    'per-query limit would be the binding constraint on coverage. In sparse '
    'categories like retail pharmacies, the true number of establishments rather than '
    'the API cap determines per-ZIP result counts.'
)
body(
    'A critical conceptual contribution of AI-generated data to commercial geography '
    'research is that its errors are analytically informative. An AI system trained '
    'on consumer-facing data sources will miss establishments with low digital '
    'visibility: recently opened businesses, independent operators with minimal web '
    'presence, and pharmacies in neighborhoods with low consumer digital engagement. '
    'When these coverage gaps are mapped spatially and compared '
    'against regulatory ground truth, the resulting pattern exposes the geography '
    'of digital invisibility, a geography that is correlated with income, race, and '
    'neighborhood disinvestment in ways that make it analytically valuable for '
    'health equity research.'
)
body(
    'The diagnostic power of AI data errors has methodological implications for '
    'commercial geography more broadly. Where NETS lag effects cause a database to '
    'overcount active pharmacies in a community experiencing chain retreat, AI '
    'consumer-visibility bias causes the opposite error, undercounting low-profile '
    'independent operators. The intersection of these two error types in the same '
    'community produces a situation in which two independent datasets provide '
    'inconsistent counts in opposite directions, and the discrepancy itself is '
    'evidence of a measurement problem that warrants attention. This is the '
    'triangulation logic motivating the three-way comparison framework developed '
    'in this study.'
)

section_heading('2.4 Ground Truth and Regulatory Data Sources')

body(
    'The selection of an appropriate ground truth dataset is the most consequential '
    'methodological decision in any validation study of commercial establishment data. '
    'Three candidates are available in the Minnesota pharmacy context. The National '
    'Council for Prescription Drug Programs DataQ database provides verified '
    'information on approximately 60,000 licensed pharmacies nationwide and is '
    'maintained by the pharmacy industry\'s own standards organization. The Minnesota '
    'Board of Pharmacy licensure database records all pharmacies currently licensed '
    'to operate in Minnesota, updated on an annual cycle, and accessible via formal '
    'data request. The NPPES NPI Registry, maintained by the Centers for Medicare '
    'and Medicaid Services, assigns National Provider Identifiers to all healthcare '
    'providers including pharmacies, but its scope extends to corporate legal '
    'entities, specialty dispensaries, and historical establishments that are not '
    'accessible to retail consumers.'
)
body(
    'This study uses NPPES as an interim ground truth pending receipt of the '
    'Minnesota Board of Pharmacy licensure dataset, which was formally requested '
    'in April 2026. NPPES has the significant practical advantage of being queryable '
    'via a public API without a data sharing agreement, enabling immediate validation '
    'while the regulatory data request is processed. However, NPPES indexing by NPI '
    'holder rather than by retail location creates a fundamental comparability problem '
    'that shapes all validation metrics reported in this study. The registry includes '
    'single NPI records held by corporate pharmacy chains rather than individual '
    'store locations, inactive establishments whose NPIs were never deactivated, and '
    'specialty dispensary operations that are not consumer-accessible retail locations.'
)
body(
    'The Minnesota Board of Pharmacy dataset is expected to provide substantially '
    'improved ground truth by restricting coverage to currently licensed retail '
    'dispensaries. Its limitations are different in character: licensure records '
    'will likely be current and retail-focused but may omit pharmacies with pending '
    'renewal applications, pharmacies that have recently relocated, and any '
    'establishments that operate under informal or unlicensed arrangements. The '
    'triangulated validation framework developed in this study, which pairs AI data '
    'against NPPES interim ground truth while treating NETS comparison as a third '
    'leg of the triangle, is designed to remain analytically informative even before '
    'the highest-quality ground truth becomes available.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: METHODOLOGY
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Chapter 3: Methodology')

body(
    'This chapter describes the research design, data collection pipeline, ground '
    'truth construction, validation framework, and spatial analysis methods used in '
    'this study. Section 3.1 defines the study area and explains the selection '
    'rationale. Section 3.2 documents the Multi-Agent System architecture and its '
    'four pipeline stages. Section 3.3 describes ground truth construction using the '
    'NPPES NPI Registry as an interim measure. Section 3.4 specifies the three '
    'pairwise comparisons in the validation framework. Section 3.5 details spatial '
    'analysis methods including the pharmacy desert classification and income '
    'quartile stratification. Section 3.6 enumerates study limitations and their '
    'mitigations.'
)

section_heading('3.1 Study Area')

body(
    'The study area is the Minneapolis-St. Paul Metropolitan Statistical Area, '
    'defined for this study by 60 ZIP codes spanning the City of Minneapolis '
    '(55401-55415, 55454, 55455), the City of St. Paul (55101-55108, 55116-55119, '
    '55130), and inner-ring suburbs including Roseville, Richfield, Edina, '
    'Bloomington, Brooklyn Park, and Coon Rapids. The spatial analysis covers 472 '
    'census tracts across Hennepin and Ramsey counties. The Twin Cities MSA was '
    'selected for four reasons.'
)
body(
    'First, the metropolitan area presents a well-documented and actively worsening '
    'pharmacy desert crisis with available longitudinal data extending to 2009, '
    'providing an unusually rich empirical baseline (Minnesota Department of Health, '
    '2024; Pereira et al., 2024). Second, the University of Minnesota College of '
    'Pharmacy maintains an active research program on pharmacy access in Minnesota, '
    'providing independent validation benchmarks and contextual knowledge of the '
    'local pharmacy landscape. Third, the Minnesota Board of Pharmacy maintains a '
    'licensure database accessible via formal data request, providing a regulatory '
    'ground truth that most states do not make available to researchers. Fourth, '
    'North Minneapolis provides a geographically contained and historically documented '
    'case study of pharmacy desert formation that tests the AI pipeline\'s sensitivity '
    'in a high-desert, low-digital-visibility urban neighborhood.'
)
body(
    'The 60 ZIP code scope was determined by the research objective of covering both '
    'core cities and the inner suburban ring. This scope captures the full range of '
    'urban commercial geography contexts relevant to pharmacy access: high-density '
    'urban cores with recent chain closures, transitional neighborhoods, and '
    'medium-density suburban areas with dispersed development patterns. Pilot data '
    'collection for the Minneapolis-only scope (17 ZIP codes) in 2025 established '
    'pipeline performance benchmarks before the full MSA deployment that produced '
    'the dataset analyzed in this thesis.'
)

section_heading('3.2 Multi-Agent System Architecture')

body(
    'The data collection and classification pipeline is implemented as a four-stage '
    'Multi-Agent System. The first stage, the Search Agent, retrieves pharmacy '
    'locations using the Google Maps Places Text Search API. To address the API\'s '
    'hard limit of 60 results per query (three pages of 20 results each), each '
    'target ZIP code is divided into a 2x2 location-biased grid. The centroid of '
    'each ZIP code is resolved using the pgeocode library, and four offset query '
    'points are derived at approximately positive and negative 780 meters latitude '
    'and 790 meters longitude from the centroid. Each grid point issues an '
    'independent location-biased search, and results are de-duplicated by Google\'s '
    'place_id identifier before further processing. This grid strategy increases '
    'recall in dense urban areas where a single centroid query might return fewer '
    'than the full establishment count within the ZIP.'
)
body(
    'The second stage, the Enrichment Agent, queries the Place Details API for each '
    'unique place_id to retrieve business name, formatted address, phone number, '
    'business status, price level, opening hours, website URL, and up to five '
    'reviews. Reviews are requested with the reviews_sort="newest" parameter to '
    'ensure that the Last_Review_Date field reflects the most recent customer '
    'activity date rather than an algorithmically ranked selection. The Wayback '
    'Machine CDX API is then queried for each business website to obtain a count '
    'of archived snapshot-years, which serves as a proxy for chain status and '
    'web longevity. Known chain pharmacies are assigned a sentinel value of negative '
    'one rather than querying Wayback, as their corporate websites provide no '
    'meaningful longevity signal for the individual store location.'
)
body(
    'The third stage, the Classification Agent, submits each enriched record to '
    'OpenAI GPT-4o-mini for NAICS classification and metadata estimation. The model '
    'receives business name, address, Google Maps type labels, opening hours, price '
    'level, and review excerpts, along with a structured classification prompt '
    'specifying the NAICS 44-45 hierarchy and distinguishing criteria for code '
    '446110 (Pharmacies and Drug Stores) relative to adjacent categories. The model '
    'outputs a six-digit NAICS code, a High or Low confidence label, a brief '
    'reasoning string, and estimated employee count and founding year. Temperature '
    'is set to 0.0 for deterministic, reproducible output. The fourth stage, the '
    'QA Agent, applies quality filters: Is_Target_Match flags records where the '
    'NAICS code matches 446110, and records with Business_Status other than '
    'OPERATIONAL are retained but flagged. Final datasets are saved as timestamped '
    'CSV files to prevent overwriting, with 22 output columns aligned to the NETS '
    'field structure.'
)
body(
    'The pipeline is implemented in Python 3.11 and orchestrated through a central '
    'workflow class that manages API rate limiting, retry logic for 429 errors, '
    'and progress checkpointing. The Google Maps client is called with exponential '
    'backoff on rate limit responses, with a maximum of three retries per query. '
    'NPPES validation uses the RapidFuzz library for fuzzy string matching, which '
    'provides the token sort ratio algorithm needed to match pharmacy names that '
    'are often abbreviated, reordered, or formatted inconsistently across data sources. '
    'The spatial analysis pipeline uses GeoPandas for all spatial join operations, '
    'with CRS management handled explicitly to prevent coordinate reference system '
    'distortions in centroid computation and distance measurement.'
)

section_heading('3.3 Ground Truth Construction')

body(
    'The National Plan and Provider Enumeration System NPI Registry is used as the '
    'primary interim ground truth. NPPES is queried at the ZIP code level using the '
    'validate_nppes.py module, which retrieves all active pharmacy NPI records '
    '(taxonomy code 333600000X) within the target ZIP codes via the CMS public API. '
    'AI-collected pharmacy names are matched to NPPES records using fuzzy string '
    'similarity with RapidFuzz token sort ratio at a threshold of 0.75. For each '
    'AI record, the highest-scoring NPPES match above threshold is recorded as a '
    'true positive. AI records with no match above threshold are false positives, '
    'and NPPES records not matched by any AI record are false negatives.'
)
body(
    'NPPES false negatives are further categorized into four classes using '
    'business name pattern matching. The closed_chain category captures records '
    'matching names of known-closed chains including Snyder Drug, Osco Drug, and '
    'Phar-Mor. The corporate_legal_name category captures records where the NPI '
    'holder is a corporate parent entity rather than an individual retail location, '
    'with SUPERVALU PHARMACIES INC as the most frequently encountered example in '
    'the Twin Cities dataset. The specialty_nonretail category captures records '
    'matching specialty pharmacy patterns including mail-order, compounding, '
    'infusion, oncology, long-term care, and hospital outpatient operations. The '
    'possible_missed_retail category contains the remaining records that may '
    'represent active retail pharmacies the AI pipeline failed to identify.'
)
body(
    'This categorization framework enables two modes of validation reporting. The '
    'raw validation metrics (Precision 38.8%, Recall 17.8%, F1 24.4%) use the '
    'full NPPES record set as the denominator and are appropriate for '
    'characterizing AI performance against the NPPES registry as administered. '
    'The adjusted recall metric uses only the possible_missed_retail category '
    'as the effective false negative count, providing a better approximation of '
    'true AI collection performance against retail-comparable ground truth. '
    'Both metrics are reported in Chapter 4, with the interpretive distinction '
    'clearly noted.'
)
body(
    'A formal data request for the Minnesota Board of Pharmacy licensure database '
    'was submitted in April 2026. Upon receipt, a dedicated validation module will '
    'be implemented with the same fuzzy matching interface as the NPPES module, '
    'enabling direct comparison of results under both ground truth standards. '
    'The Board of Pharmacy dataset is expected to substantially improve precision '
    'and recall by restricting the ground truth to currently licensed retail '
    'dispensaries and eliminating the corporate legal name and closed chain records '
    'that inflate the NPPES false negative count. NPPES validation will continue '
    'as a secondary cross-check after the regulatory data is received, enabling '
    'a three-way comparison that includes both ground truth sources simultaneously.'
)

section_heading('3.4 Validation Framework')

body(
    'The validation framework uses a triangulated comparison structure across three '
    'data sources. Comparison A, currently active, pairs the AI-collected dataset '
    'against the NPPES NPI Registry using fuzzy name matching at the ZIP code '
    'level, reporting precision, recall, F1, and false negative categorization. '
    'Comparison B, planned upon receipt of the Minnesota Board of Pharmacy dataset, '
    'will apply the same matching protocol to regulatory licensure records, '
    'providing the retail-comparable ground truth necessary for definitive '
    'assessment of AI collection performance.'
)
body(
    'Comparison C, planned for the thesis analysis phase, will spatially join '
    'AI-collected records to the NETS database by address or coordinates and '
    'compare attribute distributions for NAICS code, employee count, and '
    'establishment year. The spatial join will use a 100-meter proximity threshold '
    'consistent with address geocoding uncertainty in both datasets. The attribute '
    'comparison will test whether NETS and AI data consistently differ in their '
    'employee count estimates (reflecting the known NETS lag in recording small '
    'establishment changes) and establishment year (reflecting the NETS lag in '
    'recording business deaths). This comparison will constitute the direct '
    'validation of the thesis\'s central claim about NETS data quality in the '
    'pharmacy domain.'
)
body(
    'The triangulated approach is motivated by the recognition that no single '
    'administrative dataset constitutes a perfect ground truth. NPPES overstates '
    'pharmacy coverage through corporate legal names and stale records. NETS '
    'overstates coverage through its two-to-three-year death lag. AI data '
    'understates coverage through consumer-visibility filtering. The comparison '
    'of all three sources against each other, with the Minnesota Board of Pharmacy '
    'dataset providing the closest approximation to a retail-only ground truth, '
    'produces an analytical matrix in which discrepancies between sources carry '
    'diagnostic meaning about the construction logic of each data system.'
)

section_heading('3.5 Spatial Analysis')

body(
    'Census tract boundaries are obtained from the TIGER/Line 2023 shapefile for '
    'Hennepin and Ramsey counties, covering 472 tracts. Tract-level socioeconomic '
    'variables are retrieved from the American Community Survey five-year 2023 '
    'estimates (2019-2023 average) via the Census Data API. Key variables include '
    'median household income, total population, and percent non-White population '
    'derived from race and ethnicity tables. AI-collected pharmacy records and '
    'NPPES false negative records are spatially joined to tracts using GeoPandas '
    'sjoin with a point-in-polygon operation in EPSG:4326.'
)
body(
    'Following Qato et al. (2014), a census tract is classified as a pharmacy '
    'desert if the nearest retail pharmacy is more than 804 meters (0.5 miles) '
    'from the tract centroid. Nearest-pharmacy distance is computed using '
    'sjoin_nearest after projecting both the pharmacy point layer and tract '
    'centroids to EPSG:3857 (Web Mercator) for accurate metric distance '
    'computation. Tract centroids are computed in EPSG:3857 to avoid the '
    'coordinate reference system distortion that arises when computing centroids '
    'in geographic coordinates. All distance computations are performed in the '
    'projected coordinate system before results are converted back to geographic '
    'coordinates for mapping.'
)
body(
    'Tracts are stratified into four income quartiles based on median household '
    'income computed across all 472 tracts with non-missing income data. Desert '
    'rate (proportion of tracts classified as pharmacy deserts), pharmacy density '
    '(pharmacies per 1,000 population), and mean nearest-pharmacy distance are '
    'reported per quartile. These stratified statistics test the income-access '
    'relationship and identify whether the pharmacy desert burden is '
    'disproportionately concentrated in lower-income communities, as the health '
    'equity literature predicts (Qato et al., 2014; Dill and Gelmon, 2023). '
    'The North Minneapolis case study is treated as a separate geographic unit '
    'throughout the spatial analysis, with ZIP codes 55411 and 55412 analyzed '
    'both as part of the full MSA dataset and as a self-contained sub-region.'
)

section_heading('3.6 Limitations')

body(
    'Four limitations constrain the interpretive scope of this study. First, the '
    'AI-generated dataset is cross-sectional, capturing pharmacy locations as of '
    'April 2026 and providing no temporal dimension comparable to NETS longitudinal '
    'records. The cross-sectional design means that this study can identify the '
    'current state of pharmacy coverage as measured by each data source but cannot '
    'reconstruct the trajectory of change over time that NETS\'s multi-year '
    'architecture enables. This limitation is inherent to the data collection '
    'methodology and can only be addressed through repeated cross-sectional '
    'collection, which the pipeline is designed to support.'
)
body(
    'Second, the Google Maps Places API reflects consumer-facing digital visibility, '
    'which structurally underrepresents independent pharmacies with minimal online '
    'presence, recently opened establishments, and pharmacies in low-digital-engagement '
    'neighborhoods. This bias is directional and correlated with the health equity '
    'variables of interest in this study: neighborhoods with lower internet adoption '
    'and fewer Google Business Profile claims are likely to be lower-income and '
    'to have higher proportions of residents who are people of color. The AI '
    'dataset\'s coverage gaps therefore concentrate in precisely the communities '
    'where pharmacy access is most consequential.'
)
body(
    'Third, NPPES structural characteristics inflate the false negative count '
    'relative to what a retail-only ground truth would produce. The precision and '
    'recall figures of 38.8 percent and 17.8 percent reported in this study '
    'reflect NPPES limitations as much as AI data quality. The false negative '
    'categorization framework partially corrects for this inflation, but the '
    'possible missed retail category (n=252) still conflates genuine AI coverage '
    'gaps with NPPES records that may represent non-retail or inactive locations '
    'that pattern matching failed to identify. The Minnesota Board of Pharmacy '
    'dataset, when received, will provide the definitive retail-only denominator '
    'needed to resolve this ambiguity.'
)
body(
    'Fourth, the study covers a single metropolitan statistical area, and findings '
    'may reflect Twin Cities-specific commercial geography patterns that do not '
    'generalize to other metropolitan contexts. The MSA selection was motivated by '
    'data availability and the existing research base, not by a claim that '
    'Minneapolis-St. Paul is representative of all metropolitan pharmacy access '
    'dynamics. Multi-city replication of the pipeline across metropolitan areas '
    'with different commercial geography profiles is identified as a priority '
    'direction for future research in Chapter 6.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: RESULTS
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Chapter 4: Results')

body(
    'This chapter presents results in four sections. Section 4.1 reports AI data '
    'collection performance metrics including NAICS match rates, confidence '
    'distributions, and chain versus independent composition. Section 4.2 presents '
    'the NPPES validation results including the 717-record false negative '
    'decomposition and adjusted recall calculations. Section 4.3 reports the '
    'pharmacy desert spatial analysis results by income quartile and overall. '
    'Section 4.4 examines the North Minneapolis case study in detail, connecting '
    'specific NPPES records to documented chain closures and community access impacts.'
)

section_heading('4.1 AI Data Collection Results')

body(
    'The full MSA pharmacy collection pipeline, executed across 60 target ZIP codes '
    'using the 2x2 location-biased grid search strategy, recovered 399 unique '
    'pharmacy records across 101 distinct ZIP codes. The expansion from 60 target '
    'ZIP codes to 101 covered ZIP codes reflects the spatial extent of the grid '
    'search offsets, which captured pharmacies in adjacent ZIP codes within the '
    'offset radius even when those ZIPs were not the centroid of any query point. '
    'Of the 399 records, 396 received a NAICS code of 446110 (Pharmacies and Drug '
    'Stores), yielding a match rate of 99.2 percent. The three non-matching records '
    'received adjacent NAICS codes for medical supply and health services '
    'establishments whose Google Maps listings contained pharmacy-adjacent '
    'terminology that the classification model assigned to nearby codes.'
)
body(
    'Of all 399 records, 94.5 percent received high-confidence classifications from '
    'GPT-4o-mini. The 22 low-confidence records were concentrated among independent '
    'pharmacies with limited review data and ambiguous or abbreviated business names '
    'that provided insufficient context for the classification model to assign a '
    'NAICS code with high certainty. No low-confidence records were assigned an '
    'incorrect NAICS code at the four-digit level; all low-confidence assignments '
    'were 446110 with the model flagging uncertainty about whether the establishment '
    'was a full-service pharmacy or a limited-service dispensary. The high overall '
    'confidence rate validates the prompt engineering approach detailed in the '
    'project\'s PROMPT_GUIDE.md documentation.'
)
body(
    'Name-keyword classification identified 270 chain pharmacies and 129 independent '
    'operators among the 399 collected records. The chain category includes all CVS, '
    'Walgreens, Walmart, Target, Costco, Hy-Vee, and Cub Pharmacy locations; the '
    'independent category includes all remaining establishments whose names do not '
    'match any chain keyword pattern. The 270 to 129 chain-to-independent ratio '
    '(67.7 to 32.3 percent) reflects the Twin Cities pharmacy market structure, '
    'in which major chains dominate the commercial pharmacy landscape while '
    'independent operators maintain a significant presence, particularly in '
    'immigrant community commercial corridors and in neighborhoods where chains '
    'have reduced their footprint (Figure 1).'
)
add_figure(
    'data/figures/figure1_coverage_map.png',
    'Figure 1',
    'AI-collected pharmacy locations across the Minneapolis-St. Paul MSA '
    '(60 target ZIP codes, n=399) overlaid on census tract boundaries. '
    'Red triangles indicate NPPES possible missed retail locations (n=252).'
)
body(
    'The Wayback Machine CDX analysis provides a supplementary longevity signal '
    'for the independent pharmacy subset. Of 399 records, 195 received the sentinel '
    'value of negative one, assigned to major chain pharmacies whose corporate '
    'websites provide no meaningful longevity signal for individual store locations. '
    'One hundred records had zero Wayback snapshot-years, indicating either no '
    'website or no archived web presence. Among the 104 records with positive '
    'snapshot counts, 99 had twenty or more archived snapshot-years, indicating '
    'well-established web presence consistent with long-standing pharmacy operations. '
    'Four records had between one and seven snapshot-years, consistent with '
    'relatively recently established operations or pharmacies that joined the web '
    'relatively late in their operational history (Figure 3).'
)
add_figure(
    'data/figures/figure3_wayback_distribution.png',
    'Figure 3',
    'Wayback Machine CDX snapshot-year distribution for AI-collected pharmacies '
    '(independent operators only, n=104 with positive counts; chains excluded as '
    'sentinel -1, n=195). Bimodal distribution with peaks at zero and 20 or more '
    'snapshot-years.'
)
body(
    'The bimodal Wayback distribution, clustering at zero (no web presence) and '
    'twenty or more years (established operations), reflects the structural '
    'composition of the pharmacy sector. Pharmacies with no web presence tend '
    'to be either very recently opened establishments that have not yet established '
    'a digital footprint or long-standing independent operators in communities '
    'with limited consumer digital engagement. Pharmacies with twenty or more '
    'Wayback snapshot-years are typically independent operators who have maintained '
    'a consistent web presence since the early 2000s, indicating multi-decade '
    'operational continuity. The absence of records in the intermediate range '
    '(eight to nineteen snapshot-years) suggests limited pharmacy founding activity '
    'in the 2007-2016 period, consistent with the broader trend of pharmacy '
    'consolidation during the post-Affordable Care Act era.'
)

section_heading('4.2 Validation Against NPPES')

body(
    'The raw validation figures look poor. Precision 38.8 percent. Recall 17.8 '
    'percent. F1 24.4 percent. The NPPES NPI Registry contained 872 active pharmacy '
    'records within the 60 target ZIP codes; fuzzy name matching at a threshold of '
    '0.75 produced 155 true positives, 244 false positives, and 717 false negatives. '
    'These figures require careful interpretation, because the NPPES denominator '
    'includes corporate legal entities, inactive establishments, and specialty '
    'dispensaries that are structurally incomparable to a consumer-facing retail '
    'census.'
)
body(
    'The 244 false positives represent AI-collected records with no NPPES match '
    'above the 0.75 threshold. These records include AI-collected pharmacies whose '
    'names differ sufficiently from NPPES entries that the fuzzy match algorithm '
    'does not connect them (name format differences between the consumer-facing '
    '"CVS Pharmacy" and the NPPES-registered "CVS PHARMACY INC" can fall below '
    'threshold for certain name variants), as well as pharmacies that appear to '
    'be operating as Google Maps listings but may lack an NPI registration. The '
    'false positive count will be substantially reduced when the Minnesota Board '
    'of Pharmacy dataset is used as the ground truth, as it will include all '
    'currently licensed retail operators with their official business names.'
)
body(
    'The 717 false negatives decompose into four categories that reveal the '
    'principal structural sources of NPPES inflation relative to a retail-only '
    'ground truth. Corporate legal name records constitute the largest category '
    'at 331 records (46.2 percent of false negatives). These records represent '
    'NPI holders that are corporate parent entities rather than individual retail '
    'locations, with SUPERVALU PHARMACIES INC being the most consequential '
    'example in the Twin Cities dataset. A single corporate NPI representing '
    'the Cub Foods pharmacy network covers multiple individual Cub Pharmacy '
    'locations that the AI pipeline correctly identifies as distinct retail '
    'establishments.'
)
body(
    'The possible missed retail category contains 252 records (35.1 percent of '
    'false negatives). These records passed the pattern matching filters for '
    'closed chains, corporate legal names, and specialty non-retail operations, '
    'and thus represent the NPPES records most likely to correspond to active '
    'retail pharmacies that the AI pipeline failed to find. A subset of these '
    'records are likely true misses reflecting the consumer-visibility bias of '
    'Google Maps: independent pharmacies with no Google Business Profile, no '
    'consumer reviews, and limited digital presence that the Places API search '
    'algorithm does not return. Another subset may represent NPPES records '
    'for establishments that closed between the NPPES refresh date and the April '
    '2026 collection date and were not identified as closed chains by the '
    'pattern matching filters.'
)
body(
    'The closed chain category contains 95 records (13.2 percent). These records '
    'match names of chains that are known to have ceased operations, including '
    'Snyder Drug (a Minnesota and Wisconsin regional chain that liquidated), '
    'Osco Drug (acquired by Albertsons), and Phar-Mor. Their continued presence '
    'in NPPES as apparently active records illustrates the registry\'s lack of '
    'active deactivation procedures for closed establishments, a lag problem '
    'that parallels the NETS lag documented by Barnatchez, Crane, and Decker '
    '(2017). The specialty non-retail category contains 39 records (5.4 percent) '
    'representing compounding pharmacies, infusion services, mail-order dispensaries, '
    'and institutional pharmacies that are not accessible to retail consumers.'
)
body(
    'An adjusted recall estimate that excludes the three NPPES artifact categories '
    'from the denominator provides a better approximation of true AI collection '
    'performance. Using only the 252 possible missed retail records as the '
    'effective false negative count, the adjusted recall is 155 divided by the '
    'sum of 155 and 252, yielding 38.1 percent. This figure represents a lower '
    'bound for the recall the pipeline will achieve against the Minnesota Board '
    'of Pharmacy dataset. The raw recall of 17.8 percent and the adjusted recall '
    'of 38.1 percent together define a range within which the true collection '
    'recall likely falls, with the Board of Pharmacy validation expected to '
    'produce a result closer to the adjusted figure.'
)

section_heading('4.3 Pharmacy Desert Spatial Analysis')

body(
    'Pharmacy desert analysis using the Qato et al. (2014) half-mile threshold '
    'identified 396 of 472 analyzed census tracts as pharmacy deserts, an overall '
    'desert rate of 83.9 percent. This high overall rate reflects the geographic '
    'structure of the Twin Cities metropolitan area rather than uniformly severe '
    'access deficits. The metropolitan area is organized around two dense urban '
    'cores with high pharmacy concentrations, surrounded by inner-ring suburbs '
    'and outer suburban expanses where low population density and large tract '
    'areas place many residents farther than half a mile from any pharmacy '
    'centroid even when pharmacy services are reasonably accessible by vehicle '
    '(Figure 2a).'
)
add_figure(
    'data/figures/figure2a_desert_map.png',
    'Figure 2a',
    'Pharmacy desert classification map (Qato et al. 2014, 0.5-mile threshold) '
    'across 472 census tracts in Hennepin and Ramsey counties. Dark shading '
    'indicates desert tracts (396 of 472, overall rate 83.9%).'
)
body(
    'Income quartile stratification reveals a pattern that initially appears '
    'counterintuitive. The lowest-income quartile (Q1, n=118 tracts) has the '
    'lowest desert rate at 65.3 percent. The second quartile (Q2, n=117 tracts) '
    'has a desert rate of 83.8 percent. The third quartile (Q3, n=117 tracts) '
    'has a desert rate of 91.5 percent. The highest-income quartile (Q4, n=118 '
    'tracts) has the highest desert rate at 94.9 percent. This positive '
    'income-desert gradient reflects geographic sorting rather than a conclusion '
    'that high-income residents face worse pharmacy access than low-income '
    'residents in absolute terms.'
)
body(
    'Low-income tracts are concentrated in the urban cores of Minneapolis and '
    'St. Paul, where pharmacy density is highest relative to population and '
    'tract size, and where the half-mile centroid distance threshold is most '
    'likely to be satisfied. High-income tracts are concentrated in outer suburbs '
    'with large tract areas and dispersed residential patterns, where the straight-'
    'line centroid distance to the nearest pharmacy frequently exceeds half a mile '
    'even when a pharmacy is located within a reasonable driving distance. The '
    'centroid-based desert metric captures this geographic dispersal as a desert '
    'classification, even though many residents in these outer-suburban tracts '
    'have adequate pharmacy access by vehicle.'
)
body(
    'The income-proximity relationship, examined continuously rather than through '
    'quartile thresholds, shows a positive correlation between median household '
    'income and distance to nearest pharmacy across the 472 tracts. The correlation '
    'confirms the geographic structure explanation: wealthier tracts are on average '
    'farther from pharmacies because they are located in lower-density suburban '
    'contexts with large tract areas (Figure 2b). The policy implication differs '
    'by geography. In the urban core, particularly in North Minneapolis and the '
    'Frogtown and West Side neighborhoods of St. Paul, pharmacy deserts reflect '
    'active chain retreat from historically underserved communities. In the outer '
    'suburbs, pharmacy deserts reflect geographic dispersal in communities where '
    'vehicle ownership rates are high and the pedestrian half-mile threshold is '
    'a less appropriate access standard.'
)
add_figure(
    'data/figures/figure2b_distance_scatter.png',
    'Figure 2b',
    'Scatter plot of median household income versus distance to nearest pharmacy '
    'by census tract (n=472). Bubble size represents tract population. North '
    'Minneapolis tracts highlighted with diamond markers.'
)
body(
    'Overall pharmacy density across the 472 tracts is 0.44 pharmacies per 1,000 '
    'population, with variation by income quartile: Q1 (lowest income) has 0.50 '
    'pharmacies per 1,000 population; Q2 has 0.49; Q3 has 0.43; and Q4 (highest '
    'income) has 0.37. This density gradient is consistent with the geographic '
    'concentration of urban-core pharmacies in lower-income tracts. However, '
    'density alone does not capture access quality: a high pharmacy density in '
    'a tract with multiple nearby pharmacies provides meaningfully different '
    'access from a single pharmacy serving a very large low-income tract, even '
    'if both tracts show similar pharmacies-per-1,000 figures. The half-mile '
    'distance metric captures this spatial concentration effect in a way that '
    'simple density measures cannot.'
)

section_heading('4.4 North Minneapolis Case Study')

body(
    'One pharmacy. Two ZIP codes. Eighteen census tracts, sixteen of which qualify '
    'as pharmacy deserts.'
)
body(
    'North Minneapolis, defined by ZIP codes 55411 and 55412, is the most acute '
    'illustration of the urban-core pharmacy desert problem documented in this study. '
    'The AI collection pipeline identified one retail pharmacy across both ZIP codes: '
    'a Cub Pharmacy in ZIP 55411 with 20 Wayback snapshot-years indicating a '
    'well-established operational web presence. ZIP 55412 returned no AI-collected '
    'pharmacies. The resulting desert rate of 88.9 percent substantially exceeds '
    'the MSA-wide figure of 83.9 percent.'
)
body(
    'NPPES validation in the 55411 and 55412 ZIP codes produced five false negative '
    'records. One record, SUPERVALU PHARMACIES INC at 701 West Broadway Avenue, is '
    'classified as a corporate legal name artifact representing the Cub Foods '
    'pharmacy network parent entity. Its presence as a false negative does not '
    'indicate an AI collection failure; it indicates that the NPPES registry '
    'recorded a corporate NPI rather than the individual Cub Pharmacy location '
    'that the AI pipeline correctly identified at a nearby address.'
)
body(
    'The four remaining NPPES false negatives in North Minneapolis are classified '
    'as possible missed retail. HENNEPIN COUNTY at 2220 Plymouth Avenue North '
    '55411 likely represents the Hennepin County Human Services pharmacy serving '
    'county benefit recipients, a non-consumer-facing operation that would be '
    'appropriately classified as specialty non-retail under a stricter pattern '
    'matching rule. MERWIN BROADWAY PHARMACY at 700 West Broadway Avenue 55411 '
    'and NORTHSIDE COMMUNITY PHARMACY at 1501 Lowry Avenue North 55411 appear '
    'to be independent retail operators with limited digital visibility, consistent '
    'with the AI pipeline\'s known underrepresentation of low-profile independent '
    'pharmacies.'
)
body(
    'The fourth possible missed retail record is WALGREEN CO at 627 West Broadway '
    'Avenue 55411. That store closed in February 2023, a documented event covered '
    'by local news and confirmed by the Minnesota Department of Health. As of April '
    '2026, it remains in NPPES as an apparently active NPI holder. The AI pipeline '
    'correctly excludes it: Google Maps lists what is currently operating. NPPES '
    'records what once held an NPI. A study using NPPES as ground truth during this '
    'period would count that location as pharmacy coverage. The building is a '
    'former Walgreens. Coverage it is not.'
)
body(
    'The North Minneapolis case illustrates the convergent validity and the '
    'diagnostic limitations of the AI collection approach simultaneously. The '
    'single Cub Pharmacy returned by the Google Maps pipeline is the one retail '
    'pharmacy confirmed operational in ZIP 55411 as of April 2026. If Merwin '
    'Broadway Pharmacy and Northside Community Pharmacy are currently active, '
    'their absence from the AI dataset reflects consumer-visibility filtering '
    'rather than actual closure, and they represent genuine AI collection gaps '
    'that the Minnesota Board of Pharmacy dataset will resolve. The WALGREEN CO '
    'record, conversely, represents a genuine ground truth artifact rather than '
    'an AI gap, demonstrating that the correct interpretation of any false negative '
    'requires careful examination of the NPPES record rather than an automatic '
    'attribution of the miss to AI collection failure.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5: DISCUSSION
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Chapter 5: Discussion')

body(
    'This chapter interprets the results presented in Chapter 4 in relation to the '
    'research questions, prior literature, and methodological framework. Section 5.1 '
    'examines what the data source discrepancies reveal about the construction logic '
    'of NPPES and about AI consumer-visibility bias. Section 5.2 discusses the '
    'implications for pharmacy desert research methods and for studies that use '
    'NETS or NPPES as primary establishment inventories. Section 5.3 evaluates '
    'the portable audit pipeline as a methodological contribution, including its '
    'cost profile, reproducibility, and transferability to other metropolitan '
    'areas and business categories.'
)

section_heading('5.1 What Data Source Discrepancies Reveal')

body(
    'The 717-record false negative decomposition is the central empirical finding '
    'of this study, and its interpretation requires careful attention to what the '
    'categorization framework reveals about NPPES as a data system. The largest '
    'single category, corporate legal name records (n=331, 46.2 percent of false '
    'negatives), reflects a structural feature of how the Centers for Medicare and '
    'Medicaid Services administers the NPI program: the program assigns identifiers '
    'to NPI holders, which are often corporate entities, rather than to individual '
    'retail locations. This means that a national pharmacy chain with hundreds of '
    'Minnesota locations may hold a single NPI for the entire network, or may hold '
    'store-level NPIs for some locations and a corporate NPI for others, depending '
    'on how the chain structured its NPI registration.'
)
body(
    'The SUPERVALU PHARMACIES INC example illustrates this architecture directly. '
    'Supervalu operates in-store Cub Pharmacy counters at multiple Twin Cities '
    'Cub Foods locations. The AI pipeline identifies each Cub Pharmacy location '
    'as a distinct retail establishment, which is the correct representation for '
    'a consumer-facing dataset. NPPES records SUPERVALU PHARMACIES INC as a single '
    'corporate NPI, which is the correct representation for a healthcare provider '
    'identifier system. The mismatch between these two correct representations '
    'produces false negatives that are not attributable to AI collection failure '
    'or to NPPES data entry error but to the fundamental difference in the unit '
    'of analysis between a consumer-location dataset and a provider-identifier '
    'system.'
)
body(
    'The closed chain category (n=95, 13.2 percent) illustrates the registry\'s '
    'lag problem. Snyder Drug, Osco Drug, and Phar-Mor ceased operations years '
    'before the April 2026 data collection date. Their NPIs remain in NPPES as '
    'active records because the NPI program has no routine deactivation procedure '
    'for pharmacy closures. The AI pipeline correctly excludes them: Google Maps '
    'does not list establishments that no longer exist. These 95 records are not AI '
    'failures. They are NPPES failures that the AI pipeline inadvertently exposes.'
)
body(
    'The specialty non-retail category (n=39, 5.4 percent) captures the '
    'definitional scope difference between NPPES\'s broad taxonomy code '
    '333600000X, which encompasses all pharmacy types, and this study\'s focus '
    'on consumer-accessible retail dispensaries. Compounding pharmacies, '
    'infusion services, mail-order dispensaries, and institutional pharmacies '
    'hold valid NPIs and are legitimately included in NPPES but are not accessible '
    'to the retail consumer seeking medication pickup within walking distance. '
    'Their inclusion in the NPPES ground truth inflates the effective denominator '
    'for recall calculation without representing a retail access resource.'
)
body(
    'The AI dataset\'s directional bias runs opposite to NPPES '
    'inflation. Consumer-facing data collection via the Google Maps Places API '
    'overrepresents establishments with active Google Business Profiles, high '
    'review counts, and strong search ranking. These are predominantly chain '
    'pharmacies and well-established independent operators with consistent digital '
    'presence. Low-profile independent pharmacies, recently opened establishments, '
    'and pharmacies in neighborhoods with low consumer digital engagement are '
    'underrepresented by design. In North Minneapolis, this means the AI '
    'pipeline captures the Cub Pharmacy at West Broadway but may miss Merwin '
    'Broadway Pharmacy and Northside Community Pharmacy, both of which appear '
    'in the NPPES possible missed retail category and serve the community but '
    'have limited digital visibility.'
)

section_heading('5.2 Implications for Pharmacy Desert Research')

body(
    'The findings of this study suggest that existing pharmacy desert research '
    'conducted with either NPPES or NETS as the primary establishment inventory '
    'contain patterned measurement errors concentrated in the communities '
    'most affected by pharmacy access disparities. The directional bias of these '
    'errors is particularly concerning for policy research. NPPES overstates '
    'pharmacy coverage in communities experiencing chain retreat by retaining '
    'closed-chain records as apparently active and by including corporate legal '
    'names that count one retail network as multiple NPPES entities. NETS '
    'overstates coverage through the two-to-three-year death lag that Barnatchez, '
    'Crane, and Decker (2017) documented, which predicts that the Walgreens '
    'closure at 627 West Broadway in February 2023 would remain in NETS as an '
    'active establishment through at least 2025 or 2026.'
)
body(
    'These measurement errors are not symmetrically distributed across census '
    'tracts. The lag-induced overcount of active pharmacies is concentrated in '
    'neighborhoods experiencing active chain retreat: low-income, high-minority '
    'communities where chains have been closing at the fastest rates. The '
    'definitional overcount from corporate legal names is concentrated in '
    'communities served by Cub Foods and similar regional chain grocery-pharmacy '
    'combinations. Both sources of overcount reduce the measured pharmacy desert '
    'rate in precisely the communities where the true desert rate, based on '
    'currently operational retail locations, is highest. The practical implication '
    'is that policymakers using NPPES or NETS data to allocate pharmacy access '
    'resources may be directed away from the communities with the greatest genuine need.'
)
body(
    'The methodological implication for pharmacy desert research is that explicit '
    'validation of the establishment inventory against current regulatory ground '
    'truth is necessary before spatial analysis is conducted. The three-way '
    'comparison framework developed in this study, pairing AI-collected consumer '
    'data against NPPES interim ground truth with planned replacement by Minnesota '
    'Board of Pharmacy regulatory data, provides a model for this validation '
    'practice. The framework\'s most important feature is that it is designed to '
    'be informative even at the interim stage, before the highest-quality ground '
    'truth is available, through the false negative categorization system that '
    'separates NPPES structural artifacts from genuine collection gaps.'
)
body(
    'For the specific case of North Minneapolis, the findings support the '
    'conclusion reached by MN Department of Health (2024) and confirm its '
    'spatial distribution at the census tract level: the community faces a severe '
    'and worsening pharmacy access deficit that has been understated in analyses '
    'that retain the closed Walgreens record as evidence of pharmacy coverage. '
    'The two independent pharmacies in the NPPES possible missed retail category, '
    'Merwin Broadway Pharmacy and Northside Community Pharmacy, represent the '
    'highest-priority records for field verification in the next phase of this '
    'research: if both are currently operational, the true North Minneapolis AI '
    'recall for the possible missed retail category approaches 25 percent (one '
    'confirmed Cub Pharmacy of approximately four possible retail pharmacies), '
    'consistent with the adjusted recall figure reported for the full MSA.'
)

body(
    'These findings also speak to a broader debate in health geography about the '
    'appropriate use of commercial data sources versus regulatory data sources for '
    'spatial accessibility research. Commercial databases such as NETS and business '
    'listing services offer the advantage of broad temporal coverage, allowing '
    'longitudinal analysis of establishment entry and exit over multi-decade windows. '
    'Regulatory databases such as the Minnesota Board of Pharmacy licensure data '
    'offer the advantage of verified current operational status, but their coverage '
    'is often jurisdiction-specific and not available in a standard API format. '
    'Consumer-facing data sources such as Google Maps offer the advantage of '
    'reflecting actual consumer accessibility, but their coverage is bounded by '
    'digital engagement patterns that reliably exclude low-profile establishments '
    'in lower-income communities.'
)
body(
    'The triangulated framework developed here uses all three source types '
    'simultaneously, treating their disagreements as analytically informative rather '
    'than as errors to be resolved by privileging one source over the others. This '
    'approach is particularly well-suited to studying pharmacy access in communities '
    'experiencing rapid commercial transition, where the lag between physical closure '
    'and database update is longest and the equity consequences of measurement error '
    'are most severe. Researchers studying other categories of essential retail, '
    'including grocery stores, primary care clinics, and childcare providers, will '
    'encounter analogous measurement challenges and may benefit from applying the '
    'same triangulation logic to their specific domain.'
)

section_heading('5.3 The Portable Audit Pipeline as a Contribution')

body(
    'The Multi-Agent System developed in this thesis demonstrates that independent '
    'validation of commercial establishment databases is feasible at a cost '
    'accessible to individual researchers and small public health organizations. '
    'The full pipeline for the Minneapolis-St. Paul MSA, covering 60 ZIP codes '
    'and 399 pharmacy records, operated at a total API cost below USD 100. The '
    'Google Maps Places API charges per query, and the 2x2 grid strategy across '
    '60 ZIP codes generates approximately 240 search queries at a cost of '
    'approximately USD 50 at standard pricing. The OpenAI GPT-4o-mini '
    'classification for 399 records at temperature 0.0 costs approximately '
    'USD 5 to 10 depending on review text length. The Wayback Machine CDX API '
    'is free. The Census and NPPES APIs are free. Total marginal cost per '
    'additional city is therefore in the USD 60 to 100 range, compared to the '
    'four-figure institutional licensing fees required to access NETS.'
)
body(
    'The 99.2 percent NAICS match rate at temperature 0.0 confirms the viability '
    'of LLM-based classification for NAICS code assignment. The deterministic '
    'property is essential for reproducibility: the same record submitted to the '
    'same model with the same prompt at temperature 0.0 produces the same output '
    'on repeated runs. This property distinguishes the AI classification approach '
    'from stochastic survey or field observation methods and enables meaningful '
    'comparison of results across data collection runs that may occur days, weeks, '
    'or months apart. The reproducibility guarantee is important for longitudinal '
    'tracking: a researcher who runs the pipeline in April 2026 and again in '
    'April 2027 will receive classification results that are directly comparable '
    'for any record that appears in both runs.'
)
body(
    'The pipeline\'s transferability to other business categories was established '
    'during the pilot data collection phase, in which coffee shops, libraries, '
    'parks, gyms, grocery stores, and religious and civic organizations were '
    'collected under the same pipeline architecture using category-specific '
    'NAICS classification prompts. The pharmacy focus in this thesis was selected '
    'for the availability of regulatory ground truth and the salience of the health '
    'equity application, but the core pipeline requires only the substitution of '
    'a target category keyword and a NAICS hierarchy specification to support '
    'any other business category covered by NAICS code 44-45 (Retail Trade) or '
    'adjacent sectors. The open-source repository at github.com/McDonaldCrispyThigh/'
    'NETS-AI includes documentation for extending the pipeline to new categories '
    'and geographies.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 6: CONCLUSION
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('Chapter 6: Conclusion')

body(
    'This thesis has addressed two interconnected problems in urban health geography: '
    'the worsening pharmacy desert crisis in the Minneapolis-St. Paul metropolitan '
    'area, and the measurement infrastructure through which that crisis is identified, '
    'monitored, and acted upon. The two problems are connected because inaccurate '
    'measurement infrastructure can systematically understate the severity of the '
    'access crisis in the communities most affected by it, misdirecting both '
    'scholarly attention and policy resources. This study developed an independent '
    'measurement tool, validated it against an interim regulatory ground truth, '
    'and applied it to generate the most current census of retail pharmacy locations '
    'in the Twin Cities MSA available to date.'
)
body(
    'Three core findings emerge from the analysis. First, the AI collection pipeline '
    'achieved a 99.2 percent NAICS 446110 match rate across 399 pharmacy records, '
    'confirming that GPT-4o-mini at temperature 0.0 provides reliable deterministic '
    'pharmacy classification from structured Google Maps metadata. Second, NPPES '
    'NPI Registry validation revealed that 64.9 percent of the 717 false negatives '
    'are attributable to NPPES structural characteristics, specifically corporate '
    'legal name indexing, stale closed-chain records, and specialty non-retail '
    'entries, rather than to AI collection gaps. The adjusted recall figure of '
    '38.1 percent, compared to the raw recall of 17.8 percent, quantifies the '
    'magnitude of this structural inflation.'
)
body(
    'Third, the pharmacy desert analysis places 16 of 18 North Minneapolis census '
    'tracts in desert status. The NPPES registry, meanwhile, retains the closed '
    'Walgreens at 627 West Broadway Avenue as an apparently active record more than '
    'three years after the February 2023 closure. A policy analysis using NPPES '
    'would classify that location as a pharmacy resource. It is not. This is the '
    'practical consequence of administrative data lag, and it is not unique to NPPES.'
)
body(
    'The broader implication extends beyond the pharmacy domain. Lag and coverage '
    'gap problems are structural features of any commercial establishment database '
    'constructed from billing, credit, or administrative records rather than from '
    'current operational observation. Researchers who use these databases to document '
    'health equity outcomes must treat database currency as an active research '
    'question, not a background assumption. The triangulated validation framework '
    'developed in this study provides a practical methodology for that '
    'examination, one that can be deployed at metropolitan scale for under USD 100.'
)
body(
    'The adjusted recall figure of 38.1 percent merits interpretation in the context '
    'of what consumer-visibility bias means for equity-focused research. The AI pipeline '
    'is most likely to miss pharmacies with limited online presence, few or no Google '
    'reviews, and no claimed Google Business Profile. These characteristics cluster in '
    'independent pharmacies serving immigrant communities, which may rely on word-of-mouth '
    'and in-language community networks rather than digital platforms for customer '
    'acquisition. The structural exclusion of these pharmacies from AI-collected '
    'data is not merely a technical limitation. It reflects the same structural '
    'disadvantage that affects independent pharmacies in competition with well-resourced '
    'chains for consumer attention, financing, and eventually survival.'
)
body(
    'Three directions for future research follow directly from this study. The '
    'arrival of the Minnesota Board of Pharmacy licensure dataset will enable '
    'replacement of NPPES with a retail-comparable ground truth, providing '
    'definitive estimates of AI collection precision and recall against currently '
    'licensed dispensaries. The planned NETS comparison (Comparison C in the '
    'validation framework) will directly test whether NETS exhibits the lag and '
    'coverage gap patterns predicted by Barnatchez, Crane, and Decker (2017) in '
    'the Twin Cities pharmacy domain, providing the empirical evidence needed to '
    'evaluate the core theoretical claim of this thesis. Multi-city replication '
    'of the AI collection pipeline across metropolitan areas with different pharmacy '
    'access profiles will establish whether the North Minneapolis findings and the '
    'NPPES structural artifacts identified here reflect nationally generalizable '
    'patterns or Twin Cities-specific dynamics.'
)
body(
    'The data infrastructure supporting pharmacy desert research is itself a '
    'determinant of the health equity policies that respond to the crisis. When '
    'the databases used to identify deserts systematically undercount closures in '
    'low-income communities of color and overcount active pharmacies through '
    'corporate legal name artifacts, the resulting policy maps obscure the '
    'communities most in need of intervention. This study provides both an '
    'empirical demonstration of that obscuring effect and a reproducible, '
    'low-cost tool for correcting it. The pipeline is available, the code is '
    'open, the methodology is documented, and the validation framework is designed '
    'to improve as each additional ground truth source becomes available. What '
    'remains is the commitment to apply it wherever the data tells us to look.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('References')

refs = [
    ('Barnatchez, K., Crane, L. D., and Decker, R. (2017). An assessment of the '
     'National Establishment Time-Series (NETS) database. Finance and Economics '
     'Discussion Series 2017-110. Board of Governors of the Federal Reserve System. '
     'Washington, D.C.'),

    ('Chapple, K., and Jacobus, R. (2009). Retail trade as a route to neighborhood '
     'revitalization. In N. Pindus, H. Wial, and H. Wolman (Eds.), Urban and Regional '
     'Policy and Its Effects (Vol. 2, pp. 191-228). Brookings Institution Press.'),

    ('Dill, M. J., and Gelmon, S. (2023). Geographic access to pharmacies in the '
     'United States. Journal of the American Pharmacists Association, 63(1), 148-156. '
     'https://doi.org/10.1016/j.japh.2022.09.006'),

    ('Federal Reserve Bank of Minneapolis. (2017). An assessment of the National '
     'Establishment Time Series database. Working Paper. Minneapolis, MN.'),

    ('Fox9 Minneapolis. (2024, February 13). MN pharmacy closures accelerating, report '
     'warns of pharmacy desert increase. fox9.com.'),

    ('Meltzer, R., and Schuetz, J. (2012). Bodegas or bagel shops? Neighborhood '
     'differences in retail and household services. Economic Development Quarterly, '
     '26(1), 73-94. https://doi.org/10.1177/0891242411430328'),

    ('Minnesota Department of Health. (2024). Pharmacy Deserts in Minnesota 2009-2024. '
     'St. Paul, MN: Minnesota Department of Health.'),

    ('MinnPost. (2025, March). There is a pharmacy shortage in Minnesota. '
     'minnpost.com.'),

    ('National Community Pharmacists Association. (2026). State of the Indy Report. '
     'Alexandria, VA: NCPA. ncpanet.org.'),

    ('National Council for Prescription Drug Programs. (2024). NCPDP DataQ: '
     'Next-generation pharmacy database. Scottsdale, AZ: NCPDP. ncpdp.org.'),

    ('Neumark, D., Zhang, J., and Wall, B. (2011). Employment dynamics and business '
     'relocation: New evidence from the National Establishment Time-Series. Research '
     'in Labor Economics, 32, 1-32. https://doi.org/10.1108/S0147-9121(2011)0000032005'),

    ('Pereira, C., Tran, M., Liu, Y., and Isetts, B. (2024). The Minnesota pharmacy '
     'landscape is drying up: Pharmacy desert emergence and implications. Abstract '
     '551580, American Public Health Association Annual Meeting. Minneapolis, MN.'),

    ('PLOS ONE. (2025). Distribution of pharmacy deserts and its association with '
     'digital divide and residential redlining across United States cities. PLOS ONE, '
     '20(2), e0316789. https://doi.org/10.1371/journal.pone.0316789'),

    ('Qato, D. M., Daviglus, M. L., Wilder, J., Lee, T., Qato, D., and Lambert, B. '
     '(2014). Pharmacy deserts are prevalent in Chicago\'s predominantly minority '
     'communities, raising medication access concerns. Health Affairs, 33(8), '
     '1359-1367. https://doi.org/10.1377/hlthaff.2013.1397'),

    ('ScienceDirect. (2025). Large language models in urban planning: A systematic '
     'review and conceptual framework. Journal of Urban Technology, 32(1), 45-78.'),

    ('ScienceDirect. (2025). Predicting restaurant survival using nationwide Google '
     'Maps data. Cities, 148, 104-119.'),

    ('Shen, Y., Gao, S., Zhou, B., and Li, M. (2025). On the use of LLMs for '
     'GIS-based spatial analysis. ISPRS International Journal of Geo-Information, '
     '14(10), 401. https://doi.org/10.3390/ijgi14100401'),

    ('Star Tribune. (2024). Number of pharmacy deserts grows as chain stores close. '
     'startribune.com.'),

    ('University of Minnesota College of Pharmacy. (2024). Mapping pharmacy deserts '
     'across Minnesota. pharmacy.umn.edu.'),

    ('Walls, D. W. (2007). National Establishment Time-Series database: Data overview. '
     'SSRN Working Paper 1022962. https://doi.org/10.2139/ssrn.1022962'),

    ('Washington State Department of Commerce. (2021). Data driven insight: Evaluating '
     'the Dun and Bradstreet toolkit. Olympia, WA: Washington State Department of '
     'Commerce.'),
]

for r in refs:
    ref_entry(r)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════════════

chapter_heading('List of Figures')

for fig, caption in [
    ('Figure 1', 'AI-collected pharmacy locations across the Minneapolis-St. Paul MSA '
                 '(60 ZIP codes, n=399) overlaid on census tract boundaries. '
                 'Red triangles indicate NPPES possible missed retail locations (n=252).'),
    ('Figure 2a', 'Pharmacy desert classification map (Qato et al. 2014, 0.5-mile '
                  'threshold) across 472 census tracts in Hennepin and Ramsey counties. '
                  'Dark shading indicates desert tracts (396 of 472, overall rate 83.9%).'),
    ('Figure 2b', 'Scatter plot of median household income versus distance to nearest '
                  'pharmacy by census tract (n=472). Bubble size represents tract '
                  'population. North Minneapolis tracts highlighted with diamond markers.'),
    ('Figure 3',  'Wayback Machine CDX snapshot-year distribution for AI-collected '
                  'pharmacies (independent operators only, chains excluded as sentinel -1). '
                  'Bimodal distribution with peaks at zero and 20 or more snapshot-years.'),
]:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    _run(p, fig + '.  ', bold=True)
    _run(p, caption)


# ════════════════════════════════════════════════════════════════════════════
# SAVE AND VERIFY
# ════════════════════════════════════════════════════════════════════════════

out_path = os.path.join(os.path.dirname(__file__), 'Honors_Thesis_Zheng_2026.docx')
doc.save(out_path)
print(f'Saved: {out_path}')

from docx import Document as _D
_doc  = _D(out_path)
paras = [p.text.strip() for p in _doc.paragraphs if p.text.strip()]
em    = [p for p in paras if '\u2014' in p]
print(f'File size           : {os.path.getsize(out_path):,} bytes')
print(f'Non-empty paragraphs: {len(paras)}')
print(f'Em-dash violations  : {len(em)}')
if em:
    for v in em[:3]:
        print(f'  -> {v[:80]}')
else:
    print('Em-dash check       : PASS')
